import asyncio
import os
import re
from collections import Counter
from pathlib import Path
from typing import Any, Dict, List, Optional

import spacy

from fastapi import (
    APIRouter,
    Depends,
    File,
    Form,
    HTTPException,
    Query,
    UploadFile,
)
from fastapi.responses import FileResponse, Response
from loguru import logger
from surreal_commands import execute_command_sync, submit_command

from api.command_service import CommandService
from api.models import (
    AssetModel,
    CommonGraphCreate,
    CommonGraphResponse,
    CreateSourceInsightRequest,
    InsightCreationResponse,
    SourceCreate,
    SourceInsightResponse,
    SourceListResponse,
    SourceResponse,
    SourceStatusResponse,
    SourceUpdate,
)
from commands.source_commands import SourceProcessingInput
from open_notebook.config import UPLOADS_FOLDER
from open_notebook.database.repository import ensure_record_id, repo_query
from open_notebook.domain.common_graph import CommonGraph
from open_notebook.domain.notebook import Notebook, Source
from open_notebook.domain.transformation import Transformation
from open_notebook.exceptions import InvalidInputError, NotFoundError
from urllib.parse import unquote
from pydantic import BaseModel

router = APIRouter()


def generate_unique_filename(original_filename: str, upload_folder: str) -> str:
    """Generate unique filename like Streamlit app (append counter if file exists)."""
    file_path = Path(upload_folder)
    file_path.mkdir(parents=True, exist_ok=True)

    # Split filename and extension
    stem = Path(original_filename).stem
    suffix = Path(original_filename).suffix

    # Check if file exists and generate unique name
    counter = 0
    while True:
        if counter == 0:
            new_filename = original_filename
        else:
            new_filename = f"{stem} ({counter}){suffix}"

        full_path = file_path / new_filename
        if not full_path.exists():
            return str(full_path)
        counter += 1


async def save_uploaded_file(upload_file: UploadFile) -> str:
    """Save uploaded file to uploads folder and return file path."""
    if not upload_file.filename:
        raise ValueError("No filename provided")

    # Generate unique filename
    file_path = generate_unique_filename(upload_file.filename, UPLOADS_FOLDER)

    try:
        # Save file
        with open(file_path, "wb") as f:
            content = await upload_file.read()
            f.write(content)

        logger.info(f"Saved uploaded file to: {file_path}")
        return file_path
    except Exception as e:
        logger.error(f"Failed to save uploaded file: {e}")
        # Clean up partial file if it exists
        if os.path.exists(file_path):
            os.unlink(file_path)
        raise


def parse_source_form_data(
    type: str = Form(...),
    notebook_id: Optional[str] = Form(None),
    notebooks: Optional[str] = Form(None),  # JSON string of notebook IDs
    url: Optional[str] = Form(None),
    content: Optional[str] = Form(None),
    title: Optional[str] = Form(None),
    transformations: Optional[str] = Form(None),  # JSON string of transformation IDs
    embed: str = Form("false"),  # Accept as string, convert to bool
    delete_source: str = Form("false"),  # Accept as string, convert to bool
    async_processing: str = Form("false"),  # Accept as string, convert to bool
    file: Optional[UploadFile] = File(None),
) -> tuple[SourceCreate, Optional[UploadFile]]:
    """Parse form data into SourceCreate model and return upload file separately."""
    import json

    # Convert string booleans to actual booleans
    def str_to_bool(value: str) -> bool:
        return value.lower() in ("true", "1", "yes", "on")

    embed_bool = str_to_bool(embed)
    delete_source_bool = str_to_bool(delete_source)
    async_processing_bool = str_to_bool(async_processing)

    # Parse JSON strings
    notebooks_list = None
    if notebooks:
        try:
            notebooks_list = json.loads(notebooks)
        except json.JSONDecodeError:
            logger.error(f"Invalid JSON in notebooks field: {notebooks}")
            raise ValueError("Invalid JSON in notebooks field")

    transformations_list = []
    if transformations:
        try:
            transformations_list = json.loads(transformations)
        except json.JSONDecodeError:
            logger.error(f"Invalid JSON in transformations field: {transformations}")
            raise ValueError("Invalid JSON in transformations field")

    # Create SourceCreate instance
    try:
        source_data = SourceCreate(
            type=type,
            notebook_id=notebook_id,
            notebooks=notebooks_list,
            url=url,
            content=content,
            title=title,
            file_path=None,  # Will be set later if file is uploaded
            transformations=transformations_list,
            embed=embed_bool,
            delete_source=delete_source_bool,
            async_processing=async_processing_bool,
        )
        pass  # SourceCreate instance created successfully
    except Exception as e:
        logger.error(f"Failed to create SourceCreate instance: {e}")
        raise

    return source_data, file


COMMON_GRAPH_BLOCKLIST = {
    'said', 'told', 'asked', 'went', 'came', 'met', 'along', 'namely',
    'thereafter', 'stayed', 'rented', 'months', 'case', 'near', 'brother',
    'associates', 'car', 'village', 'person', 'people', 'man', 'men',
    'woman', 'women', 'one', 'another', 'other', 'some', 'many', 'most',
    'few', 'first', 'last', 'next', 'time', 'times', 'day', 'days', 'year',
    'years', 'back', 'good', 'new', 'still', 'really', 'very',
    'may', 'might', 'shall', 'should', 'would', 'could', 'need', 'needed',
    'needs', 'want', 'wanted', 'wants', 'get', 'gets', 'got', 'make', 'makes',
    'made', 'take', 'takes', 'took', 'taken', 'say', 'says', 'go', 'goes',
    'gone', 'see', 'sees', 'saw', 'seen', 'know', 'knows', 'knew', 'known',
    'think', 'thinks', 'thought', 'feel', 'feels', 'felt', 'look', 'looks',
    'looked', 'help', 'helps', 'helped', 'show', 'shows', 'showed', 'shown',
}

_spacy_nlp: Optional[spacy.language.Language] = None


def _get_spacy_nlp() -> Optional[spacy.language.Language]:
    global _spacy_nlp
    if _spacy_nlp is not None:
        return _spacy_nlp

    # Try transformer model first (most accurate for person NER)
    for model_name in ("en_core_web_trf", "en_core_web_lg", "en_core_web_md", "en_core_web_sm"):
    # for model_name in ("en_core_web_trf", "en_core_web_lg", "en_core_web_md", "en_core_web_sm"):
        try:
            _spacy_nlp = spacy.load(model_name)
            logger.info(f"[spaCy] Loaded model: {model_name}")
            return _spacy_nlp
        except OSError:
            continue

    logger.warning(
        "No spaCy model found. "
        "Install with: python -m spacy download en_core_web_trf"
    )
    return None


def _extract_terms_fallback(text: str) -> Counter[str]:
    cleaned = re.sub(r"[^a-z0-9\s]", " ", text.lower())
    tokens = [token for token in cleaned.split() if len(token) > 2 and token not in COMMON_GRAPH_BLOCKLIST]
    return Counter(tokens)


def _extract_terms(text: str) -> Counter[str]:
    if not text:
        return Counter()

    nlp_model = _get_spacy_nlp()
    if nlp_model is None:
        return _extract_terms_fallback(text)

    doc = nlp_model(text)
    term_freq: Counter[str] = Counter()

    for ent in doc.ents:
        if ent.label_ in {"PERSON", "GPE", "LOC", "ORG"}:
            term = ent.text.lower().strip()
            if len(term) > 2 and term not in COMMON_GRAPH_BLOCKLIST:
                term_freq[term] += 3

    for chunk in doc.noun_chunks:
        term = chunk.root.lemma_.lower().strip()
        if (
            len(term) > 2
            and not chunk.root.is_stop
            and term not in COMMON_GRAPH_BLOCKLIST
            and not term.isnumeric()
        ):
            term_freq[term] += 1

    return term_freq


def _build_common_graph_metadata(sources: List[Source]) -> Dict[str, Any]:
    term_counters = []
    source_ids = []
    for source in sources:
        text = source.full_text or source.title or ''
        term_counters.append(_extract_terms(text))
        source_ids.append(source.id or '')

    if not term_counters:
        return {}

    common_terms = set(term_counters[0].keys())
    for counter in term_counters[1:]:
        common_terms &= set(counter.keys())

    nodes: List[Dict[str, Any]] = []
    links: List[Dict[str, Any]] = []

    for idx, source in enumerate(sources):
        nodes.append(
            {
                'id': f'source:{idx}',
                'label': source.title or source.id or f'Source {idx + 1}',
                'type': 'source',
                'source_id': source.id,
            }
        )

    if not common_terms:
        return {
            'common_terms': [],
            'graph': {'nodes': nodes, 'links': []},
        }

    term_scores = {
        term: sum(counter[term] for counter in term_counters)
        for term in common_terms
    }
    selected_terms = sorted(term_scores.items(), key=lambda item: item[1], reverse=True)[:15]

    for term, score in selected_terms:
        term_id = f'term:{term}'
        nodes.append(
            {
                'id': term_id,
                'label': term,
                'type': 'term',
                'weight': score,
            }
        )
        for source_idx, counter in enumerate(term_counters):
            if counter[term] > 0:
                links.append(
                    {
                        'source': f'source:{source_idx}',
                        'target': term_id,
                        'weight': counter[term],
                    }
                )

    return {
        'common_terms': [term for term, _ in selected_terms],
        'graph': {'nodes': nodes, 'links': links},
    }


# ── Comprehensive entity extraction: persons, activities, relationships ──────
# Uses BERT (all-MiniLM-L6-v2) for semantic matching + IR patterns for extraction

_sbert_model = None


def _get_sbert_model():
    global _sbert_model
    if _sbert_model is not None:
        return _sbert_model
    try:
        from sentence_transformers import SentenceTransformer
        _sbert_model = SentenceTransformer('all-MiniLM-L6-v2')
        logger.info("[CommonGraph] Loaded all-MiniLM-L6-v2")
        return _sbert_model
    except Exception as e:
        logger.warning(f"[CommonGraph] BERT load failed: {e}")
        return None


_WORD_BLOCKLIST = {
    'name', 'names', 'full', 'alias', 'parentage', 'address', 'occupation',
    'age', 'status', 'education', 'contact', 'mobile', 'email', 'mail',
    'residence', 'resident', 'permanent', 'present', 'current', 'date',
    'sir', 'mr', 'mrs', 'ms', 'dr', 'shri', 'smt', 'km', 'late',
    'son', 'daughter', 'wife', 'husband', 'brother', 'sister',
    'father', 'mother', 'uncle', 'aunt', 'nephew', 'niece',
    'accused', 'victim', 'complainant', 'witness', 'suspect', 'informer',
    'officer', 'inspector', 'constable', 'head', 'sub', 'senior', 'junior',
    'police', 'judge', 'advocate', 'counsel', 'magistrate', 'station',
    'india', 'delhi', 'haryana', 'rajasthan', 'punjab', 'gujarat', 'mumbai',
    'uttar', 'pradesh', 'madhya', 'bihar', 'bengal', 'chennai', 'kolkata',
    'district', 'village', 'city', 'town', 'road', 'street', 'nagar',
    'colony', 'sector', 'phase', 'block', 'flat', 'house', 'near',
    'the', 'and', 'or', 'of', 'in', 'at', 'to', 'for', 'with', 'from',
    'said', 'told', 'asked', 'stated', 'mentioned', 'informed', 'reported',
    'case', 'fir', 'section', 'act', 'ipc', 'crpc', 'court',
    'unknown', 'male', 'female', 'married', 'unmarried', 'single',
    'january', 'february', 'march', 'april', 'may', 'june',
    'july', 'august', 'september', 'october', 'november', 'december',
}


def _clean_person_name(raw):
    """Clean, validate and properly space a person name."""
    import re as _re
    # Normalize whitespace
    s = _re.sub(r'\s+', ' ', raw.strip())
    # Remove trailing punctuation
    s = s.rstrip('.,;:()[]')
    # Must be 3-60 chars
    if len(s) < 3 or len(s) > 60:
        return ''
    # Must start with uppercase
    if not s[0].isupper():
        return ''
    # Must have at least one alpha char
    if not any(c.isalpha() for c in s):
        return ''
    # Check blocklist
    norm = s.lower()
    words = [w for w in norm.split() if w.isalpha()]
    if not words:
        return ''
    if all(w in _WORD_BLOCKLIST for w in words):
        return ''
    if words[0] in _WORD_BLOCKLIST:
        return ''
    return s


def _take_name_words(raw, max_words=4):
    """
    Take up to max_words words from raw that form a valid person name.
    Stops at lowercase words, blocklist words, or numbers.
    Handles '@' alias notation (e.g. 'Sandeep @ Kala Jathedi').
    """
    import re as _re
    # Split on whitespace
    words = raw.strip().split()
    clean = []
    i = 0
    while i < len(words) and len(clean) < max_words:
        w = words[i]
        # Handle @ alias separator
        if w == '@' and clean:
            clean.append(w)
            i += 1
            continue
        # Strip non-alpha from word for checking
        alpha = _re.sub(r'[^a-zA-Z]', '', w)
        if not alpha:
            break
        # Stop at blocklist
        if alpha.lower() in _WORD_BLOCKLIST:
            break
        # Stop at lowercase start (unless it's after @)
        if not alpha[0].isupper() and (not clean or clean[-1] != '@'):
            break
        clean.append(w)
        i += 1
    # Remove trailing @
    while clean and clean[-1] == '@':
        clean.pop()
    return ' '.join(clean).strip()


def _extract_all_entities(text):
    """
    Extract from IR document text:
    - persons (named individuals with roles)
    - activities (crimes, weapons, drugs, events, transactions)
    - relationships (family, associates, gang members)

    Returns:
        persons: dict {norm -> {label, role}}
        activities: list of {label, type}
        relations: list of {from, to, type, label}
    """
    import re as _re

    persons = {}
    activities = []
    relations = []
    seen_rel = set()
    seen_act = set()
    current_main = None

    lines = text.split('\n')

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # ── Structured field patterns (Name:, Accused:, etc.) ────────────
        for pat, role in [
            (r'^(?:Name|Full\s*Name)\s*[:\-]\s*(.+)', 'main'),
            (r'^Accused\s*[:\-]\s*(.+)', 'accused'),
            (r'^(?:Victim|Complainant)\s*[:\-]\s*(.+)', 'victim'),
            (r'^(?:Witness|Informer)\s*[:\-]\s*(.+)', 'witness'),
            (r'^Suspect\s*[:\-]\s*(.+)', 'suspect'),
        ]:
            m = _re.match(pat, line, _re.IGNORECASE)
            if m:
                raw = m.group(1).strip()
                name = _clean_person_name(_take_name_words(raw))
                if name:
                    norm = name.lower()
                    if norm not in persons:
                        persons[norm] = {'label': name, 'role': role}
                    if role in ('main', 'accused', 'victim', 'witness'):
                        current_main = name

        # ── Honorific prefix (Sh., Smt., Mr., etc.) ──────────────────────
        for m in _re.finditer(
            r'\b(?:Sh\.|Shri|Smt\.|Km\.|Mr\.|Mrs\.|Dr\.)\s+([A-Z][a-zA-Z\s@]{2,40})',
            line
        ):
            name = _clean_person_name(_take_name_words(m.group(1)))
            if name:
                norm = name.lower()
                if norm not in persons:
                    persons[norm] = {'label': name, 'role': 'person'}

        # ── Verb + name (arrested, nabbed, etc.) ─────────────────────────
        for m in _re.finditer(
            r'\b(?:arrested|nabbed|apprehended|detained|identified)\s+([A-Z][a-zA-Z\s@]{2,40})',
            line, _re.IGNORECASE
        ):
            name = _clean_person_name(_take_name_words(m.group(1)))
            if name:
                norm = name.lower()
                if norm not in persons:
                    persons[norm] = {'label': name, 'role': 'accused'}

        # ── Family relationship patterns ──────────────────────────────────
        for pat, rel_label in [
            (r'\bS/O\s+([A-Z][a-zA-Z\s@]{2,40})', 'father'),
            (r'\bD/O\s+([A-Z][a-zA-Z\s@]{2,40})', 'father'),
            (r'\bW/O\s+([A-Z][a-zA-Z\s@]{2,40})', 'husband'),
            (r'\bSon\s+of\s+([A-Z][a-zA-Z\s@]{2,40})', 'father'),
            (r'\bDaughter\s+of\s+([A-Z][a-zA-Z\s@]{2,40})', 'father'),
            (r'\bWife\s+of\s+([A-Z][a-zA-Z\s@]{2,40})', 'husband'),
            (r'\bBrother\s+of\s+([A-Z][a-zA-Z\s@]{2,40})', 'brother'),
            (r'\bSister\s+of\s+([A-Z][a-zA-Z\s@]{2,40})', 'sister'),
            (r'\bMother\s+of\s+([A-Z][a-zA-Z\s@]{2,40})', 'mother'),
            (r'\bFather\s+of\s+([A-Z][a-zA-Z\s@]{2,40})', 'father'),
        ]:
            for m in _re.finditer(pat, line, _re.IGNORECASE):
                rel_name = _clean_person_name(_take_name_words(m.group(1)))
                if rel_name and current_main:
                    norm_r = rel_name.lower()
                    if norm_r not in persons:
                        persons[norm_r] = {'label': rel_name, 'role': 'relative'}
                    key = (current_main.lower(), norm_r, rel_label)
                    if key not in seen_rel:
                        seen_rel.add(key)
                        relations.append({
                            'from': current_main,
                            'to': rel_name,
                            'type': 'family',
                            'label': rel_label,
                        })

        # ── Associate/gang patterns ───────────────────────────────────────
        for pat, rel_label in [
            (r'\b(?:associate|associates)\s+(?:of\s+)?([A-Z][a-zA-Z\s@]{2,40})', 'associate'),
            (r'\b(?:gang\s+member|gang\s+associate)\s+([A-Z][a-zA-Z\s@]{2,40})', 'gang member'),
            (r'\b(?:friend|close\s+friend)\s+(?:of\s+)?([A-Z][a-zA-Z\s@]{2,40})', 'friend'),
            (r'\b(?:co-accused|co\s+accused)\s+([A-Z][a-zA-Z\s@]{2,40})', 'co-accused'),
            (r'\b(?:partner|accomplice)\s+(?:of\s+)?([A-Z][a-zA-Z\s@]{2,40})', 'accomplice'),
        ]:
            for m in _re.finditer(pat, line, _re.IGNORECASE):
                assoc_name = _clean_person_name(_take_name_words(m.group(1)))
                if assoc_name and current_main:
                    norm_a = assoc_name.lower()
                    if norm_a not in persons:
                        persons[norm_a] = {'label': assoc_name, 'role': 'associate'}
                    key = (current_main.lower(), norm_a, rel_label)
                    if key not in seen_rel:
                        seen_rel.add(key)
                        relations.append({
                            'from': current_main,
                            'to': assoc_name,
                            'type': 'associate',
                            'label': rel_label,
                        })

        # ── Activity patterns ─────────────────────────────────────────────
        line_lower = line.lower()
        for keywords, act_type in [
            (['robbery', 'theft', 'murder', 'assault', 'kidnapping', 'extortion',
              'fraud', 'smuggling', 'trafficking', 'dacoity', 'rape', 'abduction',
              'cheating', 'forgery', 'bribery', 'loot', 'snatching'], 'crime'),
            (['pistol', 'revolver', 'rifle', 'gun', 'knife', 'sword', 'bomb',
              'explosive', 'weapon', 'arms', 'ammunition', 'cartridge'], 'weapon'),
            (['drugs', 'drug', 'narcotics', 'heroin', 'cocaine', 'ganja',
              'smack', 'charas', 'afeem', 'opium', 'mdma'], 'drug'),
            (['money laundering', 'hawala', 'ransom', 'extortion money',
              'cash recovery', 'illegal payment'], 'transaction'),
            (['meeting', 'encounter', 'arrest', 'raid', 'recovery', 'seizure',
              'incident', 'attack', 'firing', 'chase', 'escape'], 'event'),
        ]:
            for kw in keywords:
                if kw in line_lower and kw not in seen_act:
                    seen_act.add(kw)
                    activities.append({'label': kw, 'type': act_type})

    return persons, activities, relations


def _get_sentences(text):
    lines = []
    for line in text.split('\n'):
        line = line.strip()
        if len(line) > 20 and not line.endswith(':'):
            lines.append(line)
    return lines


def _normalize(s):
    return s.lower().strip()


def _extract_personal_details(text):
    """Extract personal details fields from IR document."""
    import re as _re
    details = {}
    field_patterns = [
        (r'^(?:Name|Full\s*Name)\s*[:\-]\s*(.+)', 'Name'),
        (r'^(?:Age|DOB|Date\s*of\s*Birth)\s*[:\-]\s*(.+)', 'Age/DOB'),
        (r'^(?:Address|Residence|Permanent\s*Address|Present\s*Address)\s*[:\-]\s*(.+)', 'Address'),
        (r'^(?:Occupation|Profession)\s*[:\-]\s*(.+)', 'Occupation'),
        (r'^(?:Education|Qualification)\s*[:\-]\s*(.+)', 'Education'),
        (r'^(?:Mobile|Contact|Phone)\s*[:\-]\s*(.+)', 'Contact'),
        (r'^(?:Email|E-mail)\s*[:\-]\s*(.+)', 'Email'),
        (r'^(?:Nationality|Religion|Caste)\s*[:\-]\s*(.+)', 'Nationality'),
        (r'^(?:Marital\s*Status|Status)\s*[:\-]\s*(.+)', 'Marital Status'),
        (r'^(?:FIR|Case\s*No|Case\s*Number)\s*[:\-]\s*(.+)', 'Case No'),
        (r'^(?:Section|Sections|Offence)\s*[:\-]\s*(.+)', 'Offence'),
        (r'^(?:Police\s*Station|PS)\s*[:\-]\s*(.+)', 'Police Station'),
        (r'^(?:District)\s*[:\-]\s*(.+)', 'District'),
        (r'^(?:State)\s*[:\-]\s*(.+)', 'State'),
        (r'^(?:Accused|Victim|Complainant|Witness)\s*[:\-]\s*(.+)', 'Role'),
    ]
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            continue
        for pat, field_name in field_patterns:
            m = _re.match(pat, line, _re.IGNORECASE)
            if m and field_name not in details:
                val = m.group(1).strip().rstrip('.,;')
                if val and len(val) < 200:
                    details[field_name] = val
    return details


def _build_personal_graph(source_idx, source_title, details):
    """Build a star graph: center=person, spokes=personal detail fields."""
    nodes = []
    links = []

    # Center node = the person
    person_name = details.get('Name', source_title.replace('.docx', '').replace('.doc', '').strip())
    center_id = f'center:{source_idx}'
    nodes.append({
        'id': center_id,
        'label': person_name,
        'type': 'person',
        'role': details.get('Role', 'main'),
        'common': True,
        'weight': 3,
    })

    # Detail nodes
    skip_fields = {'Name', 'Role'}
    for i, (field, value) in enumerate(details.items()):
        if field in skip_fields:
            continue
        nid = f'detail:{source_idx}:{i}'
        # Truncate long values
        display_val = value if len(value) <= 30 else value[:28] + '…'
        nodes.append({
            'id': nid,
            'label': display_val,
            'type': 'detail',
            'field': field,
            'weight': 1,
            'common': False,
        })
        links.append({
            'source': center_id,
            'target': nid,
            'type': 'detail',
            'label': field,
            'weight': 1,
        })

    return {'nodes': nodes, 'links': links}


def _build_family_graph(source_idx, source_title, persons, relations):
    """Build family tree graph for a source document."""
    nodes = []
    links = []
    node_ids = {}

    # Find main person
    main_person = None
    for norm, info in persons.items():
        if info['role'] in ('main', 'accused', 'victim'):
            main_person = info['label']
            break
    if not main_person and persons:
        main_person = list(persons.values())[0]['label']
    if not main_person:
        main_person = source_title.replace('.docx', '').strip()

    # Center = main person
    center_id = f'fam_center:{source_idx}'
    node_ids[main_person.lower()] = center_id
    nodes.append({
        'id': center_id,
        'label': main_person,
        'type': 'person',
        'role': 'main',
        'common': True,
        'weight': 3,
    })

    # Family relations only
    family_labels = {'father', 'mother', 'husband', 'wife', 'brother', 'sister', 'son', 'daughter', 'uncle', 'aunt'}
    for rel in relations:
        if rel['type'] != 'family':
            continue
        rel_name = rel['to']
        rel_label = rel['label']
        norm_r = rel_name.lower()
        if norm_r not in node_ids:
            nid = f'fam_rel:{source_idx}:{len(node_ids)}'
            node_ids[norm_r] = nid
            nodes.append({
                'id': nid,
                'label': rel_name,
                'type': 'relative',
                'role': rel_label,
                'common': False,
                'weight': 1,
            })
        from_id = node_ids.get(rel['from'].lower(), center_id)
        to_id = node_ids[norm_r]
        links.append({
            'source': from_id,
            'target': to_id,
            'type': 'family',
            'label': rel_label,
            'weight': 2,
        })

    # Also add relatives found in persons dict
    for norm, info in persons.items():
        if info['role'] == 'relative' and norm not in node_ids:
            nid = f'fam_rel:{source_idx}:{len(node_ids)}'
            node_ids[norm] = nid
            nodes.append({
                'id': nid,
                'label': info['label'],
                'type': 'relative',
                'role': 'relative',
                'common': False,
                'weight': 1,
            })
            links.append({
                'source': center_id,
                'target': nid,
                'type': 'family',
                'label': 'relative',
                'weight': 1,
            })

    return {'nodes': nodes, 'links': links}


def _build_associates_graph(source_idx, source_title, persons, relations, activities):
    """Build associates/gang/friends graph for a source document."""
    nodes = []
    links = []
    node_ids = {}

    # Find main person
    main_person = None
    for norm, info in persons.items():
        if info['role'] in ('main', 'accused', 'victim'):
            main_person = info['label']
            break
    if not main_person and persons:
        main_person = list(persons.values())[0]['label']
    if not main_person:
        main_person = source_title.replace('.docx', '').strip()

    center_id = f'assoc_center:{source_idx}'
    node_ids[main_person.lower()] = center_id
    nodes.append({
        'id': center_id,
        'label': main_person,
        'type': 'person',
        'role': 'main',
        'common': True,
        'weight': 3,
    })

    # Associates, gang, friends, co-accused
    assoc_types = {'associate', 'gang member', 'friend', 'co-accused', 'accomplice'}
    for rel in relations:
        if rel['type'] not in ('associate',) and rel['label'] not in assoc_types:
            continue
        assoc_name = rel['to']
        norm_a = assoc_name.lower()
        if norm_a not in node_ids:
            nid = f'assoc:{source_idx}:{len(node_ids)}'
            node_ids[norm_a] = nid
            nodes.append({
                'id': nid,
                'label': assoc_name,
                'type': 'person',
                'role': rel['label'],
                'common': False,
                'weight': 1,
            })
        from_id = node_ids.get(rel['from'].lower(), center_id)
        to_id = node_ids[norm_a]
        links.append({
            'source': from_id,
            'target': to_id,
            'type': 'associate',
            'label': rel['label'],
            'weight': 2,
        })

    # Add all accused persons as associates if no explicit relations found
    if len(nodes) <= 1:
        for norm, info in persons.items():
            if info['role'] in ('accused', 'suspect') and norm != main_person.lower():
                if norm not in node_ids:
                    nid = f'assoc:{source_idx}:{len(node_ids)}'
                    node_ids[norm] = nid
                    nodes.append({
                        'id': nid,
                        'label': info['label'],
                        'type': 'person',
                        'role': info['role'],
                        'common': False,
                        'weight': 1,
                    })
                    links.append({
                        'source': center_id,
                        'target': nid,
                        'type': 'associate',
                        'label': info['role'],
                        'weight': 1,
                    })

    # Activity nodes connected to center
    for i, act in enumerate(activities[:10]):
        nid = f'assoc_act:{source_idx}:{i}'
        nodes.append({
            'id': nid,
            'label': act['label'],
            'type': 'activity',
            'activity_type': act['type'],
            'common': False,
            'weight': 1,
        })
        links.append({
            'source': center_id,
            'target': nid,
            'type': 'activity',
            'label': act['type'],
            'weight': 1,
        })

    return {'nodes': nodes, 'links': links}


async def _build_common_graph_metadata_llm(sources, model_id, prompt=None):
    """
    Comprehensive graph: persons + activities + relationships.
    Uses docx tables directly when available for accurate person extraction.
    Also runs spaCy NER on full text to catch ALL persons mentioned.
    BERT semantic matching finds cross-doc connections.
    """
    import re as _re

    all_persons = []
    all_activities = []
    all_relations = []
    all_sentences = []
    source_main_persons = {}  # doc_idx -> {name, details}  — main person per source

    def _extract_persons_spacy(text: str) -> dict:
        """Extract ALL person names from text using spaCy NER with strict filtering."""
        nlp = _get_spacy_nlp()
        if not nlp or not text:
            return {}
        
        # Comprehensive blocklist — common words mistaken as names
        _ner_skip = {
            # Aliases/titles
            'don', 'rani', 'bua', 'dada', 'bhai', 'boss', 'madam', 'lady',
            'raja', 'seth', 'guru', 'tiger', 'lion', 'king', 'queen',
            'sir', 'mr', 'mrs', 'ms', 'dr', 'shri', 'smt', 'km',
            # Status/roles
            'unknown', 'accused', 'victim', 'witness', 'suspect',
            'male', 'female', 'person', 'individual', 'subject',
            # Places
            'india', 'delhi', 'rajasthan', 'haryana', 'punjab', 'mumbai',
            'jaipur', 'sikar', 'bikaner', 'nagaur', 'jodhpur', 'udaipur',
            # Document terms
            'part', 'section', 'report', 'case', 'police', 'court',
            'fir', 'station', 'district', 'state', 'country',
            # Common English words often capitalized
            'cash', 'period', 'time', 'date', 'year', 'month', 'day',
            'money', 'amount', 'total', 'sum', 'number', 'count',
            'first', 'second', 'third', 'last', 'next', 'previous',
            'new', 'old', 'current', 'present', 'past', 'future',
            'good', 'bad', 'best', 'worst', 'better', 'worse',
            'big', 'small', 'large', 'short', 'long', 'high', 'low',
            'yes', 'no', 'nil', 'na', 'none', 'not', 'never', 'always',
            # Gang names (not person names)
            'jathedi', 'bishnoi', 'gogi', 'banuda', 'lawrence',
            'bishnoi gang', 'jathedi gang', 'gogi gang',
            # Weapons/objects
            'revolver', 'pistol', 'gun', 'knife', 'weapon', 'arms',
            'car', 'bike', 'vehicle', 'phone', 'mobile',
            # Actions/verbs
            'arrest', 'arrested', 'caught', 'seized', 'recovered',
            'found', 'seen', 'met', 'went', 'came', 'left',
        }
        
        persons = {}
        chunk_size = 100000
        for i in range(0, min(len(text), 500000), chunk_size):
            chunk = text[i:i + chunk_size]
            try:
                doc = nlp(chunk)
                for ent in doc.ents:
                    if ent.label_ != 'PERSON':
                        continue
                    
                    raw = ent.text.strip()
                    name = _clean_person_name(_take_name_words(raw, max_words=5))
                    if not name or len(name) < 3:
                        continue
                    
                    words = name.split()
                    words_lower = [w.lower() for w in words]
                    
                    # ── Strict validation rules ──────────────────────────
                    # 1. Must have at least 2 words (single-word names are usually not real)
                    if len(words) < 2:
                        # Exception: if it's a known Indian name pattern (ends with common suffixes)
                        indian_suffixes = ('kumar', 'singh', 'sharma', 'verma', 'gupta', 'jain', 'yadav', 'choudhary', 'chaudhary')
                        if not any(words_lower[0].endswith(s) for s in indian_suffixes):
                            continue
                    
                    # 2. Skip if any word is in blocklist
                    if any(w in _ner_skip for w in words_lower):
                        continue
                    
                    # 3. Skip if all words are very short (< 4 chars each)
                    if all(len(w) < 4 for w in words):
                        continue
                    
                    # 4. Skip if contains numbers
                    if any(c.isdigit() for c in name):
                        continue
                    
                    # 5. Must start with uppercase letter
                    if not name[0].isupper():
                        continue
                    
                    # 6. Skip if it's all uppercase (likely a section header)
                    if name.isupper() and len(words) <= 2:
                        continue
                    
                    norm = name.lower()
                    if norm not in persons:
                        persons[norm] = {'label': name, 'role': 'person', 'details': '', 'source_title': ''}
            except Exception:
                continue
        return persons

    for source in sources:
        doc_idx = len(all_persons)  # current index before append
        # Try docx extraction first for accurate names + details
        file_path = source.asset.file_path if source.asset else None
        if file_path and os.path.exists(file_path) and file_path.lower().endswith(('.docx', '.doc')):
            try:
                profile = _extract_profile_from_docx(file_path)
                persons = {}
                main_name = profile.get('main_person', '')
                personal = profile.get('personal', {})

                detail_parts = []
                for k in ['Parentage', 'Date Of Birth', 'Age', 'Address', 'Occupation', 'Social Status', 'Nationality']:
                    if personal.get(k):
                        detail_parts.append(f"{k}: {personal[k]}")

                if main_name:
                    persons[main_name.lower()] = {
                        'label': main_name,
                        'role': 'main',
                        'details': ' | '.join(detail_parts),
                        'source_title': source.title or '',
                    }
                    source_main_persons[doc_idx] = {
                        'name': main_name,
                        'details': ' | '.join(detail_parts),
                    }
                for fam in profile.get('family', []):
                    n = fam['name']
                    if n and len(n) > 2:
                        persons[n.lower()] = {
                            'label': n,
                            'role': fam['relation'],
                            'details': fam.get('details', ''),
                            'source_title': source.title or '',
                        }
                for assoc in profile.get('associates', []):
                    n = assoc['name']
                    if n and len(n) > 2:
                        persons[n.lower()] = {
                            'label': n,
                            'role': assoc['relation'],
                            'details': assoc.get('details', ''),
                            'source_title': source.title or '',
                        }

                # ── Also run spaCy NER on full text to catch ALL persons ──
                text = source.full_text or ''
                spacy_persons = _extract_persons_spacy(text)
                added_from_spacy = 0
                for norm, info in spacy_persons.items():
                    if norm not in persons:
                        info['source_title'] = source.title or ''
                        persons[norm] = info
                        added_from_spacy += 1
                logger.info(f"[CommonGraph] DOCX '{source.title}': {len(persons)} persons total ({added_from_spacy} from NER)")

                all_persons.append(persons)
                _, activities, relations = _extract_all_entities(text)
                all_activities.append(activities)
                all_relations.append(relations)
                all_sentences.append(_get_sentences(text))
                continue
            except Exception as e:
                logger.warning(f"[CommonGraph] DOCX failed for '{source.title}': {e}, using text")

        text = source.full_text or source.title or ''
        persons, activities, relations = _extract_all_entities(text)
        # Also run spaCy NER for comprehensive person extraction
        spacy_persons = _extract_persons_spacy(text)
        for norm, info in spacy_persons.items():
            if norm not in persons:
                info['source_title'] = source.title or ''
                persons[norm] = info
        # Filter: keep entries that look like actual names
        valid_persons = {}
        for norm, info in persons.items():
            label = info['label']
            words = label.split()
            if not (1 <= len(words) <= 6):
                continue
            if not all(w[0].isupper() or w == '@' for w in words if w.isalpha()):
                continue
            if any(w[0].islower() and len(w) > 2 for w in words):
                continue
            if len(label) > 50:
                continue
            valid_persons[norm] = info
        all_persons.append(valid_persons)
        all_activities.append(activities)
        all_relations.append(relations)
        all_sentences.append(_get_sentences(text))
        logger.info(f"[CommonGraph] '{source.title}': {len(valid_persons)} persons (text+NER), {len(activities)} activities")

    # ── BERT semantic matching ────────────────────────────────────────────
    sbert = _get_sbert_model()
    if sbert is not None and len(sources) >= 2:
        try:
            from sentence_transformers import util as su
            for i in range(len(sources)):
                for j in range(i + 1, len(sources)):
                    si = all_sentences[i][:200]
                    sj = all_sentences[j][:200]
                    if not si or not sj:
                        continue
                    ei = sbert.encode(si, convert_to_tensor=True, show_progress_bar=False)
                    ej = sbert.encode(sj, convert_to_tensor=True, show_progress_bar=False)
                    scores = su.cos_sim(ei, ej)
                    for a in range(len(si)):
                        for b in range(len(sj)):
                            if float(scores[a][b]) > 0.72:
                                for sent, doc_idx in [(si[a], i), (sj[b], j)]:
                                    p, act, rel = _extract_all_entities(sent)
                                    for norm, info in p.items():
                                        if norm not in all_persons[doc_idx]:
                                            all_persons[doc_idx][norm] = info
                                    for a_item in act:
                                        if a_item not in all_activities[doc_idx]:
                                            all_activities[doc_idx].append(a_item)
        except Exception as e:
            logger.warning(f"[CommonGraph] BERT matching failed: {e}")

    # ── Alias/title blocklist — these are not real person names ─────────────
    _ALIAS_BLOCKLIST = {
        'don', 'rani', 'bua', 'dada', 'bhai', 'anna', 'boss', 'sir',
        'madam', 'lady', 'mama', 'chacha', 'nana', 'dadi', 'nani',
        'raja', 'seth', 'sahib', 'ji', 'baba', 'guru',
        'tiger', 'lion', 'cobra', 'panther', 'bull', 'king', 'queen',
        'master', 'chief', 'head', 'leader', 'captain', 'major',
        'inspector', 'officer', 'constable', 'judge', 'advocate',
        'unknown', 'accused', 'victim', 'witness', 'suspect',
        'male', 'female', 'person', 'individual', 'subject',
        'india', 'delhi', 'rajasthan', 'haryana', 'punjab',
        # Common IR document words mistaken as names
        'part', 'section', 'report', 'case', 'fir', 'police', 'court',
        'station', 'district', 'state', 'country', 'address', 'name',
        'age', 'date', 'birth', 'occupation', 'education', 'status',
        'nil', 'na', 'none', 'not', 'applicable',
        # Common English words that appear capitalized in documents
        'cash', 'period', 'time', 'money', 'amount', 'total', 'number',
        'first', 'second', 'third', 'last', 'next', 'new', 'old',
        'good', 'bad', 'yes', 'no', 'never', 'always', 'also',
        'year', 'month', 'day', 'week', 'hour', 'minute',
        'place', 'area', 'zone', 'region', 'location', 'spot',
        'work', 'job', 'duty', 'task', 'role', 'post',
        'gang', 'group', 'team', 'unit', 'force', 'squad',
        # Single-word aliases used in criminal records
        'revolver', 'pistol', 'gun', 'knife', 'weapon',
        'jathedi', 'bishnoi', 'gogi', 'banuda', 'lawrence',
    }

    def _is_alias_only(label: str) -> bool:
        """Return True if the label is just an alias/title, not a real name."""
        words = label.lower().split()
        # Single word that's in blocklist
        if len(words) == 1 and words[0] in _ALIAS_BLOCKLIST:
            return True
        # All words are aliases/titles
        if all(w in _ALIAS_BLOCKLIST for w in words):
            return True
        # Very short single word (< 4 chars) — likely not a real name
        if len(words) == 1 and len(words[0]) < 4:
            return True
        # Starts with a number
        if words and words[0][0].isdigit():
            return True
        # Contains numbers
        if any(c.isdigit() for c in label):
            return True
        # Single word that's a common English word (not a name)
        if len(words) == 1 and words[0].isalpha() and words[0] in {
            'cash', 'period', 'time', 'money', 'amount', 'total', 'number',
            'first', 'second', 'third', 'last', 'next', 'new', 'old',
            'year', 'month', 'day', 'place', 'area', 'work', 'job',
            'gang', 'group', 'team', 'unit', 'force', 'squad',
            'part', 'section', 'report', 'case', 'police', 'court',
        }:
            return True
        return False

    # ── Build unified maps with fuzzy name matching ───────────────────────
    person_map = {}

    def _normalize_name(s: str) -> str:
        """Normalize name for fuzzy matching — lowercase, strip spaces, remove special chars."""
        import re as _re
        s = s.lower().strip()
        # Remove common suffixes/prefixes that vary between docs
        s = _re.sub(r'\s+', ' ', s)
        # Normalize @ aliases
        s = _re.sub(r'\s*@\s*', '@', s)
        return s

    def _names_match(n1: str, n2: str) -> bool:
        """Check if two names refer to the same person (fuzzy match)."""
        a, b = _normalize_name(n1), _normalize_name(n2)
        if a == b:
            return True
        # One is a substring of the other
        if len(a) > 4 and (a in b or b in a):
            return True
        # Split by @ and check if any part matches
        a_parts = [p.strip() for p in a.split('@')]
        b_parts = [p.strip() for p in b.split('@')]
        for ap in a_parts:
            for bp in b_parts:
                if len(ap) > 4 and len(bp) > 4 and (ap == bp or ap in bp or bp in ap):
                    return True
        # First word matches (first name)
        a_words = a.split()
        b_words = b.split()
        if a_words and b_words:
            a_first = a_words[0]
            b_first = b_words[0]
            if len(a_first) > 4 and a_first == b_first:
                return True
        return False

    for doc_idx, persons in enumerate(all_persons):
        # Collect main person names for this doc to skip duplicates
        main_names_this_doc = {
            v['label'].lower() for v in persons.values() if v.get('role') == 'main'
        }

        for norm, info in persons.items():
            label = info['label']

            # Skip alias-only entries
            if _is_alias_only(label):
                continue

            # Skip if this is a partial match of the main person of this doc
            # e.g. "Anuradha" should not appear separately if "Anuradha Choudhary" is main
            is_main_partial = False
            for main_norm in main_names_this_doc:
                if norm != main_norm and _names_match(norm, main_norm):
                    is_main_partial = True
                    break
            if is_main_partial:
                continue

            # Check if this person already exists under a different spelling
            matched_key = None
            for existing_key in person_map:
                if _names_match(norm, existing_key):
                    matched_key = existing_key
                    break

            if matched_key:
                person_map[matched_key]['doc_indices'].add(doc_idx)
                existing_label = person_map[matched_key]['label']
                # Prefer the name with more words (more complete)
                if len(label.split()) > len(existing_label.split()):
                    person_map[matched_key]['label'] = label
                elif len(label.split()) == len(existing_label.split()) and len(label) > len(existing_label):
                    person_map[matched_key]['label'] = label
                # Merge details — prefer entry with more detail
                if info.get('details') and len(info.get('details', '')) > len(person_map[matched_key].get('details', '')):
                    person_map[matched_key]['details'] = info.get('details', '')
                # Prefer non-'person' role (table-extracted roles are more specific)
                if info.get('role') not in ('person', 'main') and person_map[matched_key].get('role') == 'person':
                    person_map[matched_key]['role'] = info['role']
            else:
                person_map[norm] = {
                    'label': label,
                    'role': info['role'],
                    'details': info.get('details', ''),
                    'source_title': info.get('source_title', ''),
                    'doc_indices': set(),
                }
                person_map[norm]['doc_indices'].add(doc_idx)

    logger.info(f"[CommonGraph] Person map: {len(person_map)} unique persons")
    logger.info(f"[CommonGraph] Common persons: {[v['label'] for v in person_map.values() if len(v['doc_indices']) >= 2]}")

    activity_map = {}
    for doc_idx, activities in enumerate(all_activities):
        for act in activities:
            key = act['label']
            if key not in activity_map:
                activity_map[key] = {'label': act['label'], 'type': act['type'], 'doc_indices': set()}
            activity_map[key]['doc_indices'].add(doc_idx)

    if not person_map and not activity_map:
        logger.warning("[CommonGraph] Nothing extracted, falling back to term extraction")
        return _build_common_graph_metadata(sources)

    # ── Build graph ───────────────────────────────────────────────────────
    nodes = []
    links = []
    node_ids = {}

    # Source nodes — use main person name as label (not file name)
    for idx, source in enumerate(sources):
        nid = f'source:{idx}'
        main_info = source_main_persons.get(idx, {})
        main_name_label = main_info.get('name', '') or source.title or f'Source {idx + 1}'
        nodes.append({
            'id': nid,
            'label': main_name_label,
            'type': 'source',
            'source_id': source.id,
            'details': main_info.get('details', ''),
        })

    # Person nodes — sorted: common first, then by label
    # Skip main persons (role='main') — they are represented by source nodes
    sorted_persons = sorted(
        [v for v in person_map.values() if v.get('role') != 'main'],
        key=lambda x: (-len(x['doc_indices']), x['label'])
    )
    for i, ent in enumerate(sorted_persons):
        norm = ent['label'].lower()
        nid = f'person:{i}'
        node_ids[norm] = nid
        is_common = len(ent['doc_indices']) >= 2
        nodes.append({
            'id': nid,
            'label': ent['label'],
            'type': 'person',
            'role': ent['role'],
            'details': ent.get('details', ''),
            'source_title': ent.get('source_title', ''),
            'weight': len(ent['doc_indices']),
            'common': is_common,
        })
        for src_idx in sorted(ent['doc_indices']):
            links.append({
                'source': f'source:{src_idx}',
                'target': nid,
                'type': 'appears_in',
                'weight': 1,
            })

    # Activity nodes — kept for activity graph tab
    sorted_acts = sorted(
        activity_map.values(),
        key=lambda x: (-len(x['doc_indices']), x['label'])
    )
    for i, act in enumerate(sorted_acts):
        nid = f'activity:{i}'
        node_ids[act['label']] = nid
        is_common = len(act['doc_indices']) >= 2
        nodes.append({
            'id': nid,
            'label': act['label'],
            'type': 'activity',
            'activity_type': act['type'],
            'weight': len(act['doc_indices']),
            'common': is_common,
        })
        for src_idx in sorted(act['doc_indices']):
            links.append({
                'source': f'source:{src_idx}',
                'target': nid,
                'type': 'appears_in',
                'weight': 1,
            })

    # Relationship links
    seen_links = set()
    for doc_relations in all_relations:
        for rel in doc_relations:
            from_norm = rel['from'].lower()
            to_norm = rel['to'].lower()
            from_id = node_ids.get(from_norm)
            to_id = node_ids.get(to_norm)

            if to_id is None:
                rid = f'relative:{len(node_ids)}'
                node_ids[to_norm] = rid
                to_id = rid
                nodes.append({
                    'id': rid,
                    'label': rel['to'],
                    'type': 'relative',
                    'role': rel['type'],
                    'weight': 1,
                    'common': False,
                })

            if from_id and to_id:
                link_key = f'{from_id}|{to_id}|{rel["label"]}'
                if link_key not in seen_links:
                    seen_links.add(link_key)
                    links.append({
                        'source': from_id,
                        'target': to_id,
                        'type': rel['type'],
                        'label': rel['label'],
                        'weight': 2,
                    })

    common_terms = [e['label'] for e in sorted_persons if e.get('common')]
    common_terms += [a['label'] for a in sorted_acts if a.get('common')]

    # ── Build 3 separate graphs per source ───────────────────────────────
    personal_graphs = []
    family_graphs = []
    associates_graphs = []

    for src_idx, source in enumerate(sources):
        text = source.full_text or source.title or ''
        details = _extract_personal_details(text)
        pg = _build_personal_graph(src_idx, source.title or f'Source {src_idx+1}', details)
        fg = _build_family_graph(src_idx, source.title or f'Source {src_idx+1}', all_persons[src_idx], all_relations[src_idx])
        ag = _build_associates_graph(src_idx, source.title or f'Source {src_idx+1}', all_persons[src_idx], all_relations[src_idx], all_activities[src_idx])
        personal_graphs.append({'source_title': source.title or f'Source {src_idx+1}', 'graph': pg})
        family_graphs.append({'source_title': source.title or f'Source {src_idx+1}', 'graph': fg})
        associates_graphs.append({'source_title': source.title or f'Source {src_idx+1}', 'graph': ag})

    logger.info(
        f"[CommonGraph] Final: {len(nodes)} nodes, {len(links)} links, "
        f"{sum(1 for n in nodes if n.get('common'))} common"
    )

    return {
        'common_terms': common_terms,
        'graph': {'nodes': nodes, 'links': links},
        'graph_type': 'network',
        'person_count': len(sorted_persons),
        'activity_count': len(sorted_acts),
        'common_count': sum(1 for n in nodes if n.get('common')),
        'personal_graphs': personal_graphs,
        'family_graphs': family_graphs,
        'associates_graphs': associates_graphs,
    }



@router.get("/sources", response_model=List[SourceListResponse])
async def get_sources(
    notebook_id: Optional[str] = Query(None, description="Filter by notebook ID"),
    limit: int = Query(
        50, ge=1, le=100, description="Number of sources to return (1-100)"
    ),
    offset: int = Query(0, ge=0, description="Number of sources to skip"),
    sort_by: str = Query(
        "updated", description="Field to sort by (created or updated)"
    ),
    sort_order: str = Query("desc", description="Sort order (asc or desc)"),
):
    """Get sources with pagination and sorting support."""
    try:
        # Validate sort parameters
        if sort_by not in ["created", "updated"]:
            raise HTTPException(
                status_code=400, detail="sort_by must be 'created' or 'updated'"
            )
        if sort_order.lower() not in ["asc", "desc"]:
            raise HTTPException(
                status_code=400, detail="sort_order must be 'asc' or 'desc'"
            )

        # Build ORDER BY clause
        order_clause = f"ORDER BY {sort_by} {sort_order.upper()}"

        # Build the query
        if notebook_id:
            # Verify notebook exists first
            notebook = await Notebook.get(notebook_id)
            if not notebook:
                raise HTTPException(status_code=404, detail="Notebook not found")

            # Query sources for specific notebook - include command field with FETCH
            query = f"""
                SELECT id, asset, created, title, updated, topics, command,
                (SELECT VALUE count() FROM source_insight WHERE source = $parent.id GROUP ALL)[0].count OR 0 AS insights_count,
                (SELECT VALUE id FROM source_embedding WHERE source = $parent.id LIMIT 1) != [] AS embedded
                FROM (select value in from reference where out=$notebook_id)
                {order_clause}
                LIMIT $limit START $offset
                FETCH command
            """
            result = await repo_query(
                query,
                {
                    "notebook_id": ensure_record_id(notebook_id),
                    "limit": limit,
                    "offset": offset,
                },
            )
        else:
            # Query all sources - include command field with FETCH
            query = f"""
                SELECT id, asset, created, title, updated, topics, command,
                (SELECT VALUE count() FROM source_insight WHERE source = $parent.id GROUP ALL)[0].count OR 0 AS insights_count,
                (SELECT VALUE id FROM source_embedding WHERE source = $parent.id LIMIT 1) != [] AS embedded
                FROM source
                {order_clause}
                LIMIT $limit START $offset
                FETCH command
            """
            result = await repo_query(query, {"limit": limit, "offset": offset})

        # Convert result to response model
        # Command data is already fetched via FETCH command clause
        response_list = []
        for row in result:
            command = row.get("command")
            command_id = None
            status = None
            processing_info = None

            # Extract status from fetched command object (already resolved by FETCH)
            if command and isinstance(command, dict):
                command_id = str(command.get("id")) if command.get("id") else None
                status = command.get("status")
                # Extract execution metadata from nested result structure
                result_data = command.get("result")
                execution_metadata = (
                    result_data.get("execution_metadata", {})
                    if isinstance(result_data, dict)
                    else {}
                )
                processing_info = {
                    "started_at": execution_metadata.get("started_at"),
                    "completed_at": execution_metadata.get("completed_at"),
                    "error": command.get("error_message"),
                }
            elif command:
                # Command exists but FETCH failed to resolve it (broken reference)
                command_id = str(command)
                status = "unknown"

            response_list.append(
                SourceListResponse(
                    id=row["id"],
                    title=row.get("title"),
                    topics=row.get("topics") or [],
                    asset=AssetModel(
                        file_path=row["asset"].get("file_path")
                        if row.get("asset")
                        else None,
                        url=row["asset"].get("url") if row.get("asset") else None,
                    )
                    if row.get("asset")
                    else None,
                    embedded=row.get("embedded", False),
                    embedded_chunks=0,  # Not needed in list view
                    insights_count=row.get("insights_count", 0),
                    created=str(row["created"]),
                    updated=str(row["updated"]),
                    # Status fields from fetched command
                    command_id=command_id,
                    status=status,
                    processing_info=processing_info,
                )
            )

        return response_list
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error fetching sources: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error fetching sources: {str(e)}")


@router.post("/sources", response_model=SourceResponse)
async def create_source(
    form_data: tuple[SourceCreate, Optional[UploadFile]] = Depends(
        parse_source_form_data
    ),
):
    """Create a new source with support for both JSON and multipart form data."""
    source_data, upload_file = form_data

    # Initialize file_path before try block so exception handlers can reference it
    file_path = None

    try:
        # Verify all specified notebooks exist (backward compatibility support)
        for notebook_id in source_data.notebooks or []:
            notebook = await Notebook.get(notebook_id)
            if not notebook:
                raise HTTPException(
                    status_code=404, detail=f"Notebook {notebook_id} not found"
                )

        # Handle file upload if provided
        if upload_file and source_data.type == "upload":
            try:
                file_path = await save_uploaded_file(upload_file)
            except Exception as e:
                logger.error(f"File upload failed: {e}")
                raise HTTPException(
                    status_code=400, detail=f"File upload failed: {str(e)}"
                )

        # Prepare content_state for processing
        content_state: dict[str, Any] = {}

        if source_data.type == "link":
            if not source_data.url:
                raise HTTPException(
                    status_code=400, detail="URL is required for link type"
                )
            content_state["url"] = source_data.url
        elif source_data.type == "upload":
            # Use uploaded file path or provided file_path (backward compatibility)
            final_file_path = file_path or source_data.file_path
            if not final_file_path:
                raise HTTPException(
                    status_code=400,
                    detail="File upload or file_path is required for upload type",
                )
            content_state["file_path"] = final_file_path
            content_state["delete_source"] = source_data.delete_source
        elif source_data.type == "text":
            if not source_data.content:
                raise HTTPException(
                    status_code=400, detail="Content is required for text type"
                )
            content_state["content"] = source_data.content
        else:
            raise HTTPException(
                status_code=400,
                detail="Invalid source type. Must be link, upload, or text",
            )

        # Validate transformations exist
        transformation_ids = source_data.transformations or []
        for trans_id in transformation_ids:
            transformation = await Transformation.get(trans_id)
            if not transformation:
                raise HTTPException(
                    status_code=404, detail=f"Transformation {trans_id} not found"
                )

        # Branch based on processing mode
        if source_data.async_processing:
            # ASYNC PATH: Create source record first, then queue command
            logger.info("Using async processing path")

            # Create minimal source record - let SurrealDB generate the ID
            source = Source(
                title=source_data.title or "Processing...",
                topics=[],
            )
            await source.save()

            # Add source to notebooks immediately so it appears in the UI
            # The source_graph will skip adding duplicates
            for notebook_id in source_data.notebooks or []:
                await source.add_to_notebook(notebook_id)

            try:
                # Import command modules to ensure they're registered
                import commands.source_commands  # noqa: F401

                # Submit command for background processing
                command_input = SourceProcessingInput(
                    source_id=str(source.id),
                    content_state=content_state,
                    notebook_ids=source_data.notebooks,
                    transformations=transformation_ids,
                    embed=source_data.embed,
                )

                command_id = await CommandService.submit_command_job(
                    "open_notebook",  # app name
                    "process_source",  # command name
                    command_input.model_dump(),
                )

                logger.info(f"Submitted async processing command: {command_id}")

                # Update source with command reference immediately
                # command_id already includes 'command:' prefix
                source.command = ensure_record_id(command_id)
                await source.save()

                # Return source with command info
                return SourceResponse(
                    id=source.id or "",
                    title=source.title,
                    topics=source.topics or [],
                    asset=None,  # Will be populated after processing
                    full_text=None,  # Will be populated after processing
                    embedded=False,  # Will be updated after processing
                    embedded_chunks=0,
                    created=str(source.created),
                    updated=str(source.updated),
                    command_id=command_id,
                    status="new",
                    processing_info={"async": True, "queued": True},
                )

            except Exception as e:
                logger.error(f"Failed to submit async processing command: {e}")
                # Clean up source record on command submission failure
                try:
                    await source.delete()
                except Exception:
                    pass
                # Clean up uploaded file if we created it
                if file_path and upload_file:
                    try:
                        os.unlink(file_path)
                    except Exception:
                        pass
                raise HTTPException(
                    status_code=500, detail=f"Failed to queue processing: {str(e)}"
                )

        else:
            # SYNC PATH: Execute synchronously using execute_command_sync
            logger.info("Using sync processing path")

            try:
                # Import command modules to ensure they're registered
                import commands.source_commands  # noqa: F401

                # Create source record - let SurrealDB generate the ID
                source = Source(
                    title=source_data.title or "Processing...",
                    topics=[],
                )
                await source.save()

                # Add source to notebooks immediately so it appears in the UI
                # The source_graph will skip adding duplicates
                for notebook_id in source_data.notebooks or []:
                    await source.add_to_notebook(notebook_id)

                # Execute command synchronously
                command_input = SourceProcessingInput(
                    source_id=str(source.id),
                    content_state=content_state,
                    notebook_ids=source_data.notebooks,
                    transformations=transformation_ids,
                    embed=source_data.embed,
                )

                # Run in thread pool to avoid blocking the event loop
                # execute_command_sync uses asyncio.run() internally which can't
                # be called from an already-running event loop (FastAPI)
                result = await asyncio.to_thread(
                    execute_command_sync,
                    "open_notebook",  # app name
                    "process_source",  # command name
                    command_input.model_dump(),
                    timeout=300,  # 5 minute timeout for sync processing
                )

                if not result.is_success():
                    logger.error(f"Sync processing failed: {result.error_message}")
                    # Clean up source record
                    try:
                        await source.delete()
                    except Exception:
                        pass
                    # Clean up uploaded file if we created it
                    if file_path and upload_file:
                        try:
                            os.unlink(file_path)
                        except Exception:
                            pass
                    raise HTTPException(
                        status_code=500,
                        detail=f"Processing failed: {result.error_message}",
                    )

                # Get the processed source
                if not source.id:
                    raise HTTPException(status_code=500, detail="Source ID is missing")
                processed_source = await Source.get(source.id)
                if not processed_source:
                    raise HTTPException(
                        status_code=500, detail="Processed source not found"
                    )

                embedded_chunks = await processed_source.get_embedded_chunks()
                return SourceResponse(
                    id=processed_source.id or "",
                    title=processed_source.title,
                    topics=processed_source.topics or [],
                    asset=AssetModel(
                        file_path=processed_source.asset.file_path
                        if processed_source.asset
                        else None,
                        url=processed_source.asset.url
                        if processed_source.asset
                        else None,
                    )
                    if processed_source.asset
                    else None,
                    full_text=processed_source.full_text,
                    embedded=embedded_chunks > 0,
                    embedded_chunks=embedded_chunks,
                    created=str(processed_source.created),
                    updated=str(processed_source.updated),
                    # No command_id or status for sync processing (legacy behavior)
                )

            except Exception as e:
                logger.error(f"Sync processing failed: {e}")
                # Clean up uploaded file if we created it
                if file_path and upload_file:
                    try:
                        os.unlink(file_path)
                    except Exception:
                        pass
                raise

    except HTTPException:
        # Clean up uploaded file on HTTP exceptions if we created it
        if file_path and upload_file:
            try:
                os.unlink(file_path)
            except Exception:
                pass
        raise
    except InvalidInputError as e:
        # Clean up uploaded file on validation errors if we created it
        if file_path and upload_file:
            try:
                os.unlink(file_path)
            except Exception:
                pass
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Error creating source: {str(e)}")
        # Clean up uploaded file on unexpected errors if we created it
        if file_path and upload_file:
            try:
                os.unlink(file_path)
            except Exception:
                pass
        raise HTTPException(status_code=500, detail=f"Error creating source: {str(e)}")


@router.post("/sources/json", response_model=SourceResponse)
async def create_source_json(source_data: SourceCreate):
    """Create a new source using JSON payload (legacy endpoint for backward compatibility)."""
    # Convert to form data format and call main endpoint
    form_data = (source_data, None)
    return await create_source(form_data)


@router.post(
    "/sources/common-graphs",
    response_model=CommonGraphResponse,
    status_code=201,
)
async def create_common_graph(request: CommonGraphCreate):
    """Create a persistent common graph from selected sources."""
    try:
        if not request.source_ids or len(request.source_ids) < 2:
            raise HTTPException(
                status_code=400,
                detail="At least two source IDs are required to create a common graph.",
            )

        source_objects: List[Source] = []
        for source_id in request.source_ids:
            try:
                source = await Source.get(source_id)
            except NotFoundError:
                raise HTTPException(
                    status_code=404,
                    detail=f"Source {source_id} not found",
                )
            source_objects.append(source)

        graph_metadata = await _build_common_graph_metadata_llm(
            source_objects,
            model_id=request.model_id,
            prompt=request.prompt,
        ) if request.model_id else _build_common_graph_metadata(source_objects)
        metadata = {"created_from": "common_graph_ui", **graph_metadata}

        common_graph = CommonGraph(
            title=request.title,
            source_ids=request.source_ids,
            status="completed",
            metadata=metadata,
        )
        await common_graph.save()

        return CommonGraphResponse(
            id=common_graph.id or "",
            title=common_graph.title,
            source_ids=common_graph.source_ids,
            status=common_graph.status,
            metadata=common_graph.metadata,
            created=str(common_graph.created),
            updated=str(common_graph.updated),
        )
    except HTTPException:
        raise
    except InvalidInputError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Error creating common graph: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Error creating common graph: {str(e)}",
        )


@router.get("/sources/common-graphs/{common_graph_id}", response_model=CommonGraphResponse)
async def get_common_graph(common_graph_id: str):
    """Retrieve a previously saved common graph."""
    try:
        common_graph = await CommonGraph.get(common_graph_id)
        return CommonGraphResponse(
            id=common_graph.id or "",
            title=common_graph.title,
            source_ids=common_graph.source_ids,
            status=common_graph.status,
            metadata=common_graph.metadata,
            created=str(common_graph.created),
            updated=str(common_graph.updated),
        )
    except NotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except Exception as e:
        logger.error(f"Error fetching common graph {common_graph_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error fetching common graph: {str(e)}")


async def _resolve_source_file(source_id: str) -> tuple[str, str]:
    source = await Source.get(source_id)
    if not source:
        raise HTTPException(status_code=404, detail="Source not found")

    file_path = source.asset.file_path if source.asset else None
    if not file_path:
        raise HTTPException(status_code=404, detail="Source has no file to download")

    safe_root = os.path.realpath(UPLOADS_FOLDER)
    resolved_path = os.path.realpath(file_path)

    if not resolved_path.startswith(safe_root):
        logger.warning(
            f"Blocked download outside uploads directory for source {source_id}: {resolved_path}"
        )
        raise HTTPException(status_code=403, detail="Access to file denied")

    if not os.path.exists(resolved_path):
        raise HTTPException(status_code=404, detail="File not found on server")

    filename = os.path.basename(resolved_path)
    return resolved_path, filename


def _is_source_file_available(source: Source) -> Optional[bool]:
    if not source or not source.asset or not source.asset.file_path:
        return None

    file_path = source.asset.file_path
    safe_root = os.path.realpath(UPLOADS_FOLDER)
    resolved_path = os.path.realpath(file_path)

    if not resolved_path.startswith(safe_root):
        return False

    return os.path.exists(resolved_path)


@router.get("/sources/{source_id}/profile-image")
async def get_source_profile_image(source_id: str):
    """Extract the front-facing person photo from a docx source file using face detection."""
    try:
        source = await Source.get(source_id)
        if not source:
            raise HTTPException(status_code=404, detail="Source not found")

        file_path = source.asset.file_path if source.asset else None
        if not file_path or not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="No file available")

        if not file_path.lower().endswith(('.docx', '.doc')):
            raise HTTPException(status_code=404, detail="Not a docx file")

        try:
            from docx import Document as DocxDocument
            import io
            doc = DocxDocument(file_path)

            # Collect all candidate images > 5KB
            candidates = []
            for rel in doc.part.rels.values():
                if 'image' in rel.reltype:
                    try:
                        img_data = rel.target_part.blob
                        if len(img_data) < 5000:
                            continue  # skip tiny icons/logos
                        candidates.append(img_data)
                    except Exception:
                        continue

            if not candidates:
                raise HTTPException(status_code=404, detail="No image found in document")

            logger.info(f"[ProfileImage] {len(candidates)} candidate images for {source_id}")

            # ── Use OpenCV face detection to find front-facing photo ──────
            chosen = None
            try:
                import cv2
                import numpy as np

                cascade_path = cv2.data.haarcascades + 'haarcascade_frontalface_default.xml'
                face_cascade = cv2.CascadeClassifier(cascade_path)

                best_face_img = None
                best_face_area = 0

                for img_data in candidates:
                    try:
                        nparr = np.frombuffer(img_data, np.uint8)
                        img_cv = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
                        if img_cv is None:
                            continue
                        gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
                        faces = face_cascade.detectMultiScale(
                            gray,
                            scaleFactor=1.1,
                            minNeighbors=4,
                            minSize=(40, 40),
                        )
                        if len(faces) > 0:
                            # Pick image with largest detected face
                            max_area = max(int(w) * int(h) for (x, y, w, h) in faces)
                            if max_area > best_face_area:
                                best_face_area = max_area
                                best_face_img = img_data
                    except Exception:
                        continue

                if best_face_img is not None:
                    chosen = best_face_img
                    logger.info(f"[ProfileImage] Face detected, area={best_face_area}px²")
                else:
                    logger.info(f"[ProfileImage] No face detected, falling back to first portrait")

            except Exception as e:
                logger.warning(f"[ProfileImage] Face detection failed: {e}")

            # ── Fallback: first portrait-shaped image ─────────────────────
            if chosen is None:
                try:
                    from PIL import Image as PILImage
                    for img_data in candidates:
                        try:
                            img = PILImage.open(io.BytesIO(img_data))
                            w, h = img.size
                            if h / w >= 0.8 and w >= 80 and h >= 80:
                                chosen = img_data
                                break
                        except Exception:
                            continue
                except ImportError:
                    pass

            # ── Final fallback: first image ───────────────────────────────
            if chosen is None:
                chosen = candidates[0]

            content_type = 'image/jpeg'
            if chosen[:4] == b'\x89PNG':
                content_type = 'image/png'
            elif chosen[:4] == b'GIF8':
                content_type = 'image/gif'
            logger.info(f"[ProfileImage] Returning {len(chosen)} bytes, type={content_type}")
            return Response(content=chosen, media_type=content_type)

        except HTTPException:
            raise
        except Exception as e:
            logger.warning(f"Could not extract image from docx: {e}")

        raise HTTPException(status_code=404, detail="No image found in document")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


def extract_part_iv_structured(text: str) -> dict:
    """
    Fallback: extract PART IV sections from plain text using regex.
    Used only when no DOCX file is available.
    """
    text = re.sub(r'[–—]', '-', text)
    pattern = r'PART\s*-\s*IV(.*?)(?=PART\s*-\s*[VIX]+|\Z)'
    match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
    if not match:
        return {}
    content = match.group(1).strip()
    content = re.sub(r'^POINTS?\s+FOR\s+FOLLOW\s*[-–]?\s*UP\s*:[-\s]*', '', content, flags=re.IGNORECASE).strip()
    sections = {}
    parts = re.split(r'(?:^|\n)\s*([A-Za-z][A-Za-z\s\-/()]{2,50}:-)', content, flags=re.MULTILINE)
    current_key = None
    for part in parts:
        part = part.strip()
        if not part:
            continue
        if part.endswith(":-"):
            current_key = part.replace(":-", "").strip()
            sections[current_key] = ""
        elif current_key:
            sections[current_key] += part + "\n"
    return {k: v.strip() for k, v in sections.items() if v.strip()}


def _extract_part_iv_from_docx(file_path: str) -> dict:
    """
    Primary extractor: reads the DOCX directly with python-docx.

    Strategy:
    - Find the paragraph whose text matches PART IV (any dash variant).
    - Collect all following paragraphs until the next PART heading or end.
    - Detect sub-section headings by bold runs (e.g. "Education:-", "How he got involved in crime:-").
    - Handles both:
        a) Standalone heading paragraphs  ("How he got involved in crime:-")
        b) Inline heading + content on same line  ("Education:- 12th passed from...")
    - Strip the "Points for follow-up:" header paragraph.
    - Return {'sections': {heading: body_text}, 'raw': full_narrative, 'found': True}
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    doc = DocxDocument(file_path)
    paragraphs = doc.paragraphs

    # ── Step 1: locate PART IV paragraph ─────────────────────────────────
    part4_idx = None
    for i, para in enumerate(paragraphs):
        if re.search(r'PART\s*[–—\-]\s*IV\b', para.text, re.IGNORECASE):
            part4_idx = i
            break

    if part4_idx is None:
        return {'sections': {}, 'raw': '', 'found': False}

    # ── Step 2: collect paragraphs until next PART heading ────────────────
    body_paras = []
    for para in paragraphs[part4_idx + 1:]:
        txt = para.text.strip()
        if not txt:
            continue
        if re.match(r'^PART\s*[–—\-]\s*[VIX\d]', txt, re.IGNORECASE):
            break
        body_paras.append(para)

    if not body_paras:
        return {'sections': {}, 'raw': '', 'found': True}

    # ── Step 3: helpers ───────────────────────────────────────────────────
    def _para_text(para) -> str:
        return re.sub(r'\s+', ' ', para.text).strip()

    def _split_inline_heading(para):
        """
        If a paragraph starts with bold runs that end with ':-',
        split it into (heading_label, remaining_content).
        Returns (None, full_text) if no inline heading found.
        """
        runs = para.runs
        heading_parts = []
        content_start = 0
        for idx, run in enumerate(runs):
            if run.bold:
                heading_parts.append(run.text)
                content_start = idx + 1
            else:
                break  # first non-bold run ends the heading

        heading_raw = ''.join(heading_parts).strip()
        # Must end with :- to be a heading
        if heading_raw.endswith(':-') or heading_raw.endswith(': -'):
            label = re.sub(r'\s*:[-\s]*$', '', heading_raw).strip()
            # Remaining content = all runs after the bold heading runs
            remaining = ''.join(r.text for r in runs[content_start:]).strip()
            remaining = re.sub(r'\s+', ' ', remaining).strip()
            return label, remaining
        return None, _para_text(para)

    def _is_pure_heading(para) -> bool:
        """True if the entire paragraph is a heading (all bold, ends with :-)."""
        txt = para.text.strip()
        if not txt:
            return False
        non_empty = [r for r in para.runs if r.text.strip()]
        if not non_empty:
            return False
        all_bold = all(r.bold for r in non_empty)
        ends_with_colon = txt.endswith(':-') or txt.endswith(': -')
        # Pure heading: all bold AND ends with :- AND no sentence content after
        if all_bold and ends_with_colon:
            return True
        # Or: all bold, short, no sentence punctuation (e.g. "POINTS FOR FOLLOW-UP:")
        if all_bold and len(txt) < 80 and not re.search(r'[.!?]\s+[A-Z]', txt):
            return True
        return False

    # ── Step 4: build sections dict and raw narrative ─────────────────────
    sections: dict = {}
    current_key: str | None = None
    raw_lines: list[str] = []

    # Skip the "POINTS FOR FOLLOW-UP:" header paragraph
    start_idx = 0
    if body_paras and re.match(r'^POINTS?\s+FOR\s+FOLLOW', body_paras[0].text, re.IGNORECASE):
        start_idx = 1

    for para in body_paras[start_idx:]:
        txt = _para_text(para)
        if not txt:
            continue

        if _is_pure_heading(para):
            # Standalone heading paragraph
            label = re.sub(r'\s*:[-\s]*$', '', txt).strip()
            current_key = label
            if current_key not in sections:
                sections[current_key] = ''
        else:
            # Check for inline heading (bold prefix + content on same line)
            label, content = _split_inline_heading(para)
            if label:
                current_key = label
                if current_key not in sections:
                    sections[current_key] = ''
                if content:
                    sections[current_key] = (sections[current_key] + '\n' + content).strip()
                    raw_lines.append(content)
            else:
                # Pure content paragraph
                raw_lines.append(txt)
                if current_key is not None:
                    sections[current_key] = (sections[current_key] + '\n' + txt).strip()

    # Clean up section values
    sections = {k: v.strip() for k, v in sections.items() if v.strip()}

    # Raw = all non-heading content paragraphs joined with blank lines
    raw = '\n\n'.join(raw_lines)

    return {'sections': sections, 'raw': raw, 'found': True}


@router.get("/sources/{source_id}/part-iv")
async def get_source_part_iv(source_id: str):
    """Extract and return structured PART IV content from a source document."""
    try:
        source = await Source.get(source_id)
        if not source:
            raise HTTPException(status_code=404, detail="Source not found")

        # ── Primary path: read DOCX directly (most accurate) ─────────────
        file_path = source.asset.file_path if source.asset else None
        if file_path and os.path.exists(file_path) and file_path.lower().endswith(('.docx', '.doc')):
            try:
                result = _extract_part_iv_from_docx(file_path)
                result['source_id'] = source_id
                logger.info(
                    f"[PartIV] DOCX extraction: found={result['found']}, "
                    f"sections={list(result.get('sections', {}).keys())}, "
                    f"raw_len={len(result.get('raw', ''))}"
                )
                return result
            except Exception as e:
                logger.warning(f"[PartIV] DOCX extraction failed: {e}, falling back to text")

        # ── Fallback: regex on full_text ──────────────────────────────────
        text = source.full_text or ''
        if not text.strip():
            return {'sections': {}, 'raw': '', 'source_id': source_id, 'found': False}

        sections = extract_part_iv_structured(text)

        normalized = re.sub(r'[–—]', '-', text)
        pattern = r'PART\s*-\s*IV(.*?)(?=PART\s*-\s*[VIX]+|\Z)'
        match = re.search(pattern, normalized, re.IGNORECASE | re.DOTALL)
        raw_content = match.group(1).strip() if match else ''
        raw_content = re.sub(
            r'^POINTS?\s+FOR\s+FOLLOW\s*[-–]?\s*UP\s*:[-\s]*', '',
            raw_content, flags=re.IGNORECASE
        ).strip()

        return {
            'sections': sections,
            'raw': raw_content,
            'source_id': source_id,
            'found': bool(match),
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error extracting Part IV for source {source_id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/sources/{source_id}/word-cloud")
async def get_source_word_cloud(source_id: str):
    """Extract word frequency data for word cloud visualization."""
    try:
        source = await Source.get(source_id)
        if not source:
            raise HTTPException(status_code=404, detail="Source not found")

        text = source.full_text or source.title or ''
        words = _extract_word_cloud_data(text)
        return {'words': words, 'source_id': source_id}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error extracting word cloud: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/sources/{source_id}/person-context")
async def get_person_context(source_id: str, name: str = Query(..., description="Person name to find context for")):
    """Find paragraphs in the source text that mention the given person name."""
    try:
        source = await Source.get(source_id)
        if not source:
            raise HTTPException(status_code=404, detail="Source not found")

        text = source.full_text or ''
        if not text or not name:
            return {'paragraphs': [], 'name': name}

        import re as _re

        # Split into paragraphs (by double newline or single newline for short lines)
        paragraphs = [p.strip() for p in _re.split(r'\n{2,}', text) if p.strip()]
        if len(paragraphs) <= 3:
            # Fallback: split by single newline
            paragraphs = [p.strip() for p in text.split('\n') if p.strip()]

        # Find paragraphs containing the name (case-insensitive, word boundary)
        name_parts = name.lower().split()
        # Match if any significant part of the name appears
        search_terms = [p for p in name_parts if len(p) > 3]
        if not search_terms:
            search_terms = name_parts

        matching = []
        for para in paragraphs:
            para_lower = para.lower()
            # Check if any search term appears in this paragraph
            if any(term in para_lower for term in search_terms):
                # Skip very short paragraphs (headers, labels)
                if len(para) < 20:
                    continue
                # Skip if it's just a field label (e.g. "Name: Ankit")
                if _re.match(r'^[A-Za-z\s/]+:\s*\S', para) and len(para) < 60:
                    continue
                matching.append(para)

        # Return up to 5 most relevant paragraphs
        # Prioritize paragraphs where the full name appears
        full_name_lower = name.lower()
        def score(p: str) -> int:
            pl = p.lower()
            s = sum(1 for t in search_terms if t in pl)
            if full_name_lower in pl:
                s += 10
            return s

        matching.sort(key=score, reverse=True)
        return {'paragraphs': matching[:5], 'name': name}

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting person context: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# def _extract_word_cloud_data(text: str) -> list:
#     """Extract word frequencies for word cloud, filtering stopwords."""
#     import re as _re
#     from collections import Counter

#     # Comprehensive stopwords
#     stopwords = {
#         'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
#         'of', 'with', 'by', 'from', 'is', 'was', 'are', 'were', 'be', 'been',
#         'being', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would',
#         'could', 'should', 'may', 'might', 'shall', 'can', 'not', 'no', 'nor',
#         'so', 'yet', 'both', 'either', 'neither', 'each', 'few', 'more', 'most',
#         'other', 'some', 'such', 'than', 'too', 'very', 'just', 'also', 'as',
#         'if', 'then', 'that', 'this', 'these', 'those', 'it', 'its', 'he', 'she',
#         'they', 'we', 'you', 'i', 'me', 'him', 'her', 'them', 'us', 'his', 'their',
#         'our', 'your', 'my', 'who', 'which', 'what', 'when', 'where', 'how', 'why',
#         'all', 'any', 'both', 'each', 'every', 'into', 'through', 'during', 'before',
#         'after', 'above', 'below', 'between', 'out', 'off', 'over', 'under', 'again',
#         'further', 'then', 'once', 'here', 'there', 'about', 'against', 'along',
#         'around', 'because', 'while', 'although', 'however', 'therefore', 'thus',
#         'hence', 'since', 'until', 'unless', 'whether', 'though', 'even', 'only',
#         'also', 'back', 'up', 'down', 'said', 'told', 'asked', 'went', 'came',
#         'got', 'get', 'go', 'come', 'take', 'make', 'know', 'see', 'look', 'use',
#         'find', 'give', 'think', 'tell', 'become', 'show', 'leave', 'feel', 'put',
#         'bring', 'begin', 'keep', 'hold', 'write', 'stand', 'hear', 'let', 'mean',
#         'set', 'meet', 'run', 'pay', 'sit', 'speak', 'lie', 'lead', 'read', 'grow',
#         'lose', 'fall', 'send', 'build', 'stay', 'reach', 'kill', 'remain', 'suggest',
#         'raise', 'pass', 'sell', 'require', 'report', 'decide', 'pull', 'nil', 'na',
#         'not', 'applicable', 'available', 'pending', 'trial', 'investigation', 'case',
#         'status', 'accused', 'bailed', 'court', 'yet', 'arrested', 'police', 'station',
#     }

#     # Clean text
#     cleaned = _re.sub(r'[^a-zA-Z\s]', ' ', text.lower())
#     words = cleaned.split()

#     # Filter: length > 3, not stopword, not purely numeric
#     filtered = [
#         w for w in words
#         if len(w) > 3
#         and w not in stopwords
#         and not w.isnumeric()
#         and not all(c in 'ivxlcdm' for c in w)  # skip roman numerals
#     ]

#     # Count frequencies
#     counter = Counter(filtered)

#     # Return top 80 words with frequency
#     result = [
#         {'text': word, 'value': count}
#         for word, count in counter.most_common(80)
#         if count >= 2  # only words appearing 2+ times
#     ]

#     return result

def _extract_word_cloud_data(text: str) -> list:
    """Extract word frequencies for word cloud using advanced NLP (spaCy with phrases + entities)."""
    import re as _re
    from collections import Counter
    import spacy

    # Load spaCy model
    nlp = spacy.load("en_core_web_sm")

    # Clean text (preserve your logic)
    cleaned = _re.sub(r'\s+', ' ', text)

    doc = nlp(cleaned)

    tokens = []

    # 1️⃣ Single meaningful words (same as before but improved)
    for token in doc:
        if (
            token.is_alpha
            and not token.is_stop
            and not token.is_punct
            and len(token.text) > 3
            and token.pos_ in {"NOUN", "PROPN", "ADJ"}
        ):
            tokens.append(token.lemma_.lower())

    # 2️⃣ Noun phrases (KEY improvement)
    for chunk in doc.noun_chunks:
        phrase = chunk.text.lower().strip()

        # filter small/irrelevant chunks
        if len(phrase) > 4 and not any(t.is_stop for t in chunk):
            tokens.append(phrase)

    # 3️⃣ Named entities (boost important real-world info)
    for ent in doc.ents:
        if ent.label_ in {"PERSON", "ORG", "GPE", "LOC", "LAW"}:
            entity = ent.text.lower().strip()
            if len(entity) > 3:
                tokens.append(entity)

    # Count frequencies
    counter = Counter(tokens)

    # Return top 80 with frequency threshold
    result = [
        {'text': word, 'value': count}
        for word, count in counter.most_common(80)
        if count >= 2
    ]
    print(result)
    return result


@router.get("/sources/{source_id}/profile-graph")
async def get_source_profile_graph(source_id: str, model_id: Optional[str] = Query(None)):
    """Extract personal details, family, and associates from a source document."""
    try:
        source = await Source.get(source_id)
        if not source:
            raise HTTPException(status_code=404, detail="Source not found")

        # Try to read docx directly for best accuracy
        file_path = source.asset.file_path if source.asset else None
        if file_path and os.path.exists(file_path) and file_path.lower().endswith(('.docx', '.doc')):
            try:
                result = _extract_profile_from_docx(file_path)
                result['source_id'] = source_id
                result['source_title'] = source.title or 'Unknown'
                logger.info(f"[ProfileGraph] DOCX: {len(result.get('personal',{}))} personal, {len(result.get('family',[]))} family, {len(result.get('associates',[]))} associates")
                print(result)
                return result
            except Exception as e:
                logger.warning(f"[ProfileGraph] DOCX extraction failed: {e}, falling back to text")

        text = source.full_text or source.title or ''
        logger.info(f"[ProfileGraph] Text extraction, length: {len(text)}")

        # Try LLM if model available
        if model_id:
            try:
                result = await _extract_profile_graph_llm(text, model_id)
                result['source_id'] = source_id
                result['source_title'] = source.title or 'Unknown'
                return result
            except Exception as e:
                logger.warning(f"LLM profile extraction failed: {e}, falling back to regex")

        result = _extract_profile_graph(text)
        result['source_id'] = source_id
        result['source_title'] = source.title or 'Unknown'
        return result
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error extracting profile graph: {e}")
        raise HTTPException(status_code=500, detail=str(e))


def _extract_profile_from_docx(file_path: str) -> dict:
    """
    Extract profile data directly from docx tables using python-docx.

    Table layout in IR documents:
      Table 0  : Cover / interrogation report (skip)
      Table 1  : Personal details — first cell is 'name'
      Tables 2+ with first_cell == 'relation' : Family members
      Tables with first_cell == 'nationality' : Associates / gang members

    All tables are processed regardless of index so documents with varying
    numbers of family/associate tables are handled correctly.
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    doc = DocxDocument(file_path)
    personal: dict = {}
    family: list = []
    associates: list = []
    main_person: str = ''

    # ── helpers ───────────────────────────────────────────────────────────
    def _clean(s: str) -> str:
        import re as _re
        s = _re.sub(r'\s+', ' ', s.strip())
        return s.rstrip('.,;:()[]').strip()

    def _guess_gender(name: str, relation: str) -> str:
        female_rel = {
            'mother', 'wife', 'sister', 'daughter', 'aunt', 'niece',
            'girlfriend', 'bua', 'mausi', 'nani', 'dadi', 'bhabhi',
        }
        male_rel = {
            'father', 'husband', 'brother', 'son', 'uncle', 'nephew',
            'chacha', 'mama', 'nana', 'dada', 'jija', 'sala',
        }
        rl = relation.lower()
        if any(r in rl for r in female_rel):
            return 'female'
        if any(r in rl for r in male_rel):
            return 'male'
        female_sfx = ('a', 'i', 'devi', 'bai', 'kumari', 'rani', 'priya', 'lata', 'vati', 'wati')
        if any(name.lower().endswith(s) for s in female_sfx):
            return 'female'
        return 'male'

    # Values that mean "no data"
    _EMPTY_VALS = {'nil', 'n/a', 'na', 'not applicable', 'not available',
                   '-', '', 'none', 'n.a.', 'n.a', 'not known', 'unknown'}

    def _is_empty(val: str) -> bool:
        return val.lower().strip() in _EMPTY_VALS

    # Fields to skip from personal details
    _SKIP_KEYS = {
        'previous involvements', 'fir no', 'police station', 'status of case',
        'status of accused', 'action taken', 'source country', 'route of smuggling',
        'carrier', 'recipient', 'repayment', 'visit to india', 'circumstances',
        'case registered', 'tattoo image', 'deformity image', 'interrogation report',
        'network details', 'details of hide outs',
    }

    def _should_skip_personal(key: str, val: str) -> bool:
        kl = key.lower().strip()
        if any(s in kl for s in _SKIP_KEYS):
            return True
        return _is_empty(val)

    # Relations that belong to family (not associates)
    _FAMILY_RELATIONS = {
        'father', 'mother', 'brother', 'sister', 'son', 'daughter',
        'wife', 'husband', 'spouse', 'uncle', 'aunt', 'bua', 'mausi',
        'chacha', 'chachi', 'mama', 'mami', 'nana', 'nani', 'dada', 'dadi',
        'jija', 'sala', 'sali', 'bhabhi', 'devar', 'nanad',
    }

    def _is_family_relation(relation: str) -> bool:
        rl = relation.lower().strip()
        return any(fr in rl for fr in _FAMILY_RELATIONS)

    def _row_to_dict(table) -> dict:
        """Convert a table's rows into a {field_lower: value} dict."""
        d: dict = {}
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) < 2:
                continue
            key = cells[0].strip()
            # Take last non-empty cell value (tables repeat across columns)
            val = next((c for c in reversed(cells[1:]) if c.strip()), '')
            val = _clean(val)
            if key and val and not _is_empty(val):
                d[key.lower().strip()] = val
        return d

    # ── Process each table ────────────────────────────────────────────────
    for table in doc.tables:
        rows = table.rows
        if not rows:
            continue

        first_cell = rows[0].cells[0].text.strip().lower() if rows[0].cells else ''

        # ── Personal details table (first cell == 'name') ─────────────────
        if first_cell == 'name':
            for row in rows:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) < 2:
                    continue
                key = cells[0].strip()
                val = next((c for c in reversed(cells[1:]) if c.strip()), '')
                val = _clean(val)
                if not key or not val or _should_skip_personal(key, val):
                    continue
                key_title = key.title().strip()
                if key_title not in personal:
                    personal[key_title] = val
                if key.lower() in ('name', 'full name') and not main_person:
                    main_person = val
            continue

        # ── Family / relative tables (first cell == 'relation') ───────────
        if first_cell == 'relation':
            d = _row_to_dict(table)
            if not d:
                continue

            relation = d.get('relation', 'relative')
            # Name & Age field (common IR format: "Anar Singh, age 48 years")
            name_age = d.get('name & age', d.get('name', ''))
            name = name_age.split(',')[0].strip() if name_age else ''
            age_raw = name_age.split(',')[1].strip() if ',' in name_age else ''

            if not name or len(name) < 2:
                continue

            # Build details string
            details_parts: list[str] = []
            if age_raw:
                details_parts.append(f"Age: {age_raw}")
            if d.get('parentage'):
                details_parts.append(f"Parentage: {d['parentage']}")
            if d.get('address'):
                details_parts.append(f"Address: {d['address']}")
            if d.get('occupation'):
                details_parts.append(f"Occupation: {d['occupation']}")
            if d.get('contact no.') or d.get('contact'):
                details_parts.append(f"Contact: {d.get('contact no.') or d.get('contact', '')}")
            if d.get('education'):
                details_parts.append(f"Education: {d['education']}")

            entry = {
                'name': name,
                'relation': relation.lower().strip(),
                'gender': _guess_gender(name, relation),
                'details': ' | '.join(details_parts),
            }

            if _is_family_relation(relation):
                family.append(entry)
            else:
                associates.append(entry)
            continue

        # ── Associate tables (first cell == 'nationality') ────────────────
        if first_cell == 'nationality':
            d = _row_to_dict(table)
            if not d:
                continue

            name_age = d.get('name & age', d.get('name', ''))
            name = name_age.split(',')[0].strip() if name_age else ''
            age_raw = name_age.split(',')[1].strip() if ',' in name_age else ''

            if not name or len(name) < 2:
                continue

            relation = d.get('description', d.get('relation', 'associate'))
            if not relation or _is_empty(relation):
                relation = 'associate'

            # Build rich details string
            details_parts = []
            if age_raw:
                details_parts.append(f"Age: {age_raw}")
            if d.get('parentage'):
                details_parts.append(f"S/O {d['parentage']}")
            if d.get('address'):
                details_parts.append(f"Address: {d['address']}")
            if d.get('occupation'):
                details_parts.append(f"Occupation: {d['occupation']}")
            if d.get('nationality') and not _is_empty(d['nationality']):
                details_parts.append(f"Nationality: {d['nationality']}")
            if d.get('fir no') and not _is_empty(d['fir no']):
                details_parts.append(f"FIR: {d['fir no']}")
            if d.get('police station') and not _is_empty(d['police station']):
                details_parts.append(f"PS: {d['police station']}")
            if d.get('status of accused') and not _is_empty(d['status of accused']):
                details_parts.append(f"Status: {d['status of accused']}")

            associates.append({
                'name': name,
                'relation': relation.lower().strip(),
                'gender': _guess_gender(name, relation),
                'details': ' | '.join(details_parts),
            })
            continue

        # ── Catch-all: any table with 'name' or 'name & age' field ──────────
        # Handles IR docs where associate tables have different first cells
        d_all = _row_to_dict(table)
        name_age_val = d_all.get('name & age', d_all.get('name', ''))
        if name_age_val and len(name_age_val) > 2:
            name = name_age_val.split(',')[0].strip()
            age_raw = name_age_val.split(',')[1].strip() if ',' in name_age_val else ''
            if name and len(name) > 2 and name[0].isupper():
                relation = d_all.get('relation', d_all.get('description', 'associate'))
                if not relation or _is_empty(relation):
                    relation = 'associate'
                details_parts = []
                if age_raw:
                    details_parts.append(f"Age: {age_raw}")
                if d_all.get('parentage'):
                    details_parts.append(f"S/O {d_all['parentage']}")
                if d_all.get('address'):
                    details_parts.append(f"Address: {d_all['address']}")
                if d_all.get('occupation'):
                    details_parts.append(f"Occupation: {d_all['occupation']}")
                entry = {
                    'name': name,
                    'relation': relation.lower().strip(),
                    'gender': _guess_gender(name, relation),
                    'details': ' | '.join(details_parts),
                }
                if _is_family_relation(relation):
                    if not any(f['name'].lower() == name.lower() for f in family):
                        family.append(entry)
                else:
                    if not any(a['name'].lower() == name.lower() for a in associates):
                        associates.append(entry)

    logger.info(
        f"[ProfileGraph] DOCX extracted: {len(personal)} personal, "
        f"{len(family)} family, {len(associates)} associates"
    )

    # ── Also extract PART-III names from paragraphs ───────────────────────
    # These are gang/associate names listed as plain text after "Associates/Groups"
    import re as _re
    in_associates_section = False
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        tl = text.lower()
        # Detect section headers — enter associates section on PART-III
        if _re.search(r'part\s*[-–\s]*iii\b', tl) or 'associates/groups' in tl or 'gangster' in tl or 'friends/associates' in tl:
            in_associates_section = True
            continue
        # Exit associates section on PART-IV or any other PART (I, II, V, etc.)
        if _re.search(r'part\s*[-–\s]*(?:iv|i\b|ii\b|v\b|vi\b|4\b|1\b|2\b)', tl) or 'points for follow' in tl:
            in_associates_section = False
            continue
        if in_associates_section:
            # Skip section headers and PART markers
            if _re.search(r'^part\s*[-–\s]*[ivxIVX\d]+', tl):
                continue
            # Skip lines that are clearly not names
            if any(x in tl for x in ['anti gang', 'places', 'location', 'nil', 'n/a', 'besides', 'anti-gangster', 'district', 'state', 'associates/groups', 'gangster', 'friends/associates', 'gang unit', 'anti-gang']):
                continue
            # Must be short (names are short)
            if len(text) > 80:
                continue
            # Must look like a name: 1-6 words
            words = text.split()
            if not words or len(words) > 6:
                continue
            # First word must start with capital letter (not a number or symbol)
            if not words[0][0].isupper():
                continue
            # Skip pure section/header lines (all caps, short)
            if text.isupper() and len(words) <= 3:
                continue
            # Skip if it looks like a sentence (has lowercase words in middle)
            if any(w[0].islower() and len(w) > 3 and w not in ('and', 'or', 'of', '@') for w in words[1:]):
                continue
            # Skip if contains digits that look like dates/numbers (not names)
            if _re.search(r'\b\d{4}\b', text):
                continue
            name = _clean(text)
            if name and len(name) > 2:
                # Check not already in associates
                if not any(a['name'].lower() == name.lower() for a in associates):
                    associates.append({
                        'name': name,
                        'relation': 'gang associate',
                        'gender': _guess_gender(name, 'associate'),
                        'details': '',
                    })

    logger.info(f"[ProfileGraph] After PART-III: {len(associates)} associates total")

    return {
        'personal': personal,
        'family': family,
        'associates': associates,
        'main_person': main_person or personal.get('Name', ''),
    }




async def _extract_profile_graph_llm(text: str, model_id: str) -> dict:
    """
    Use LLM to dynamically extract ALL fields present in the document.
    Returns structured data without hardcoded field names.
    """
    from langchain_core.messages import HumanMessage, SystemMessage
    from open_notebook.ai.provision import provision_langchain_model
    import json as _json
    import re as _re

    system_prompt = """You are an expert at extracting structured information from Indian Police Investigation Reports (IR) and case documents.

Analyze the document carefully and extract EVERY piece of personal information about the PRIMARY SUBJECT (main accused/victim).

Return a JSON object with exactly this structure:
{
  "main_person": "full name of the primary subject",
  "personal": {
    "Name": "full name",
    "Alias": "alias or nick name if any",
    "Age": "age",
    "Date of Birth": "DOB",
    "Gender": "Male/Female",
    "Parentage": "father's name (S/O or D/O)",
    "Address": "full address",
    "Occupation": "job/work",
    "Education": "qualification",
    "Mobile": "phone number",
    "Marital Status": "married/unmarried",
    "Complexion": "fair/dark/wheatish",
    "Height": "height",
    "Eyes": "eye color",
    "Hair": "hair color/type",
    "Build": "body build",
    "Mark of Identification": "any marks/moles/tattoos",
    "Facebook ID": "facebook profile if mentioned",
    "Criminal Record": "previous cases if any",
    "Case No": "FIR/case number",
    "Sections": "IPC sections",
    "Police Station": "PS name",
    "District": "district",
    "State": "state",
    "Role": "accused/victim/witness"
  },
  "family": [...],
  "associates": [...]
}

IMPORTANT:
- Only include fields that are ACTUALLY PRESENT in the document
- Extract the EXACT values as written in the document
- Do NOT include fields that are not in the document — omit them entirely
- Do NOT use null, None, or empty string values — skip missing fields
- For personal dict: use the field names exactly as shown above
- Return ONLY the JSON object, no explanation, no markdown"""

    truncated = text[:6000]
    payload = [
        SystemMessage(content=system_prompt),
        HumanMessage(content=truncated),
    ]

    chain = await provision_langchain_model(str(payload), model_id, "transformation", max_tokens=2048)
    response = await chain.ainvoke(payload)
    raw = str(response.content if hasattr(response, 'content') else response).strip()
    logger.info(f"[ProfileGraph] LLM raw response: {raw}")

    # Strip markdown fences
    if '```' in raw:
        raw = _re.sub(r'```(?:json)?\s*', '', raw).strip()

    # Extract JSON object
    start = raw.find('{')
    end = raw.rfind('}')
    if start == -1 or end == -1:
        raise ValueError(f"No JSON object in LLM response: {raw[:200]}")

    json_str = raw[start:end + 1]

    # Fix common LLM JSON issues:
    # 1. Replace Python None with null
    json_str = json_str.replace(': None', ': null').replace(':None', ':null')
    # 2. Replace trailing commas before } or ]
    json_str = _re.sub(r',\s*([}\]])', r'\1', json_str)
    # 3. Replace single quotes with double quotes (carefully)
    # Only do this if standard parse fails
    try:
        data = _json.loads(json_str)
    except _json.JSONDecodeError as e:
        logger.warning(f"[ProfileGraph] JSON parse error: {e}, trying cleanup")
        # Remove null values that might be causing issues
        json_str = _re.sub(r'"[^"]+"\s*:\s*null\s*,?\s*', '', json_str)
        json_str = _re.sub(r',\s*([}\]])', r'\1', json_str)
        data = _json.loads(json_str)

    # Normalize
    personal = data.get('personal', {})
    if not isinstance(personal, dict):
        personal = {}
    # Ensure all values are strings (LLM may return numbers, lists, etc.)
    personal = {
        str(k): str(v) if not isinstance(v, (list, dict)) else ', '.join(str(i) for i in v) if isinstance(v, list) else str(v)
        for k, v in personal.items()
        if v is not None and str(v).strip()
    }

    family = []
    for p in data.get('family', []):
        if isinstance(p, dict) and p.get('name'):
            family.append({
                'name': str(p.get('name', '')).strip(),
                'relation': str(p.get('relation', 'relative')).strip(),
                'gender': str(p.get('gender', 'male')).strip(),
                'details': str(p.get('details', '')).strip(),
            })

    associates = []
    for p in data.get('associates', []):
        if isinstance(p, dict) and p.get('name'):
            associates.append({
                'name': str(p.get('name', '')).strip(),
                'relation': str(p.get('relation', 'associate')).strip(),
                'gender': str(p.get('gender', 'male')).strip(),
                'details': str(p.get('details', '')).strip(),
            })

    logger.info(f"[ProfileGraph] LLM extracted: {len(personal)} fields, {len(family)} family, {len(associates)} associates")

    return {
        'main_person': str(data.get('main_person', '')).strip(),
        'personal': personal,
        'family': family,
        'associates': associates,
    }


def _extract_profile_graph(text: str) -> dict:
    """
    Comprehensive extraction from IR documents (table-based docx format).
    Handles: Field | Value table rows, section-based content, narrative text.
    """
    import re as _re

    personal = {}
    family = []
    associates = []
    seen = set()

    def _guess_gender(name: str, relation: str) -> str:
        female_rel = {'mother', 'wife', 'sister', 'daughter', 'aunt', 'niece', 'girlfriend', 'smt', 'km', 'beti', 'behen', 'mata'}
        male_rel = {'father', 'husband', 'brother', 'son', 'uncle', 'nephew', 'boyfriend', 'sh', 'shri', 'beta', 'bhai', 'pita', 'parentage'}
        rl = relation.lower()
        if any(r in rl for r in female_rel): return 'female'
        if any(r in rl for r in male_rel): return 'male'
        female_sfx = ('a', 'i', 'devi', 'bai', 'kumari', 'rani', 'priya', 'lata', 'vati', 'wati')
        if any(name.lower().endswith(s) for s in female_sfx): return 'female'
        return 'male'

    def _clean(s: str) -> str:
        s = _re.sub(r'\s+', ' ', s.strip())
        return s.rstrip('.,;:()[]').strip()

    def _take_name(raw: str) -> str:
        words = raw.strip().split()[:5]
        clean = []
        bl = {'unknown', 'nil', 'n/a', 'na', 'not', 'available', 'mentioned', 'none', '-', 'r/o', 's/o', 'd/o', 'w/o'}
        for w in words:
            alpha = _re.sub(r'[^a-zA-Z@]', '', w)
            if not alpha or alpha.lower() in bl: break
            if not (alpha[0].isupper() or alpha == '@'): break
            clean.append(w)
        while clean and clean[-1] == '@': clean.pop()
        return ' '.join(clean).strip()

    def _add_family(name: str, relation: str):
        n = _take_name(name)
        if n and n.lower() not in seen and len(n) > 2:
            seen.add(n.lower())
            family.append({'name': n, 'relation': relation, 'gender': _guess_gender(n, relation), 'details': ''})

    def _add_associate(name: str, relation: str, details: str = ''):
        n = _take_name(name)
        if n and n.lower() not in seen and len(n) > 2:
            seen.add(n.lower())
            associates.append({'name': n, 'relation': relation, 'gender': _guess_gender(n, relation), 'details': details})

    # ── Step 1: Parse table rows (Field | Value format from docx) ─────────
    # The docx tables get extracted as "Field | Value | Value | Value" lines
    # or as "Field\nValue" pairs
    lines = text.split('\n')
    current_section = ''

    # Fields to skip (non-personal)
    skip_fields = {
        'note', 'remarks', 'description', 'summary', 'information',
        'subject', 'report', 'reference', 'sr no', 'serial', 'sl no', 'page',
        'previous involvements', 'status of case', 'status of accused',
        'fir no', 'police station', 'action taken', 'source country',
        'route of smuggling', 'carrier', 'recipient', 'repayment',
        'visit to india', 'circumstances', 'case registered',
        'interrogation report', 'tattoo image', 'deformity image',
    }

    # Family field names
    family_fields = {
        'parentage': 'father', 'father': 'father', 'mother': 'mother',
        'wife': 'wife', 'husband': 'husband', 'brother': 'brother',
        'sister': 'sister', 'son': 'son', 'daughter': 'daughter',
        'spouse': 'spouse',
    }

    # Associate/friend field names
    associate_fields = {
        'details of close friends during studies': 'friend',
        'close friends': 'friend',
        'friends': 'friend',
        'associates': 'associate',
        'known associates': 'associate',
        'gang members': 'gang member',
    }

    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            continue

        # Detect section headers
        line_lower = line_stripped.lower()
        if 'part-ii' in line_lower or 'family member' in line_lower:
            current_section = 'family'
            continue
        elif 'part-iii' in line_lower or 'associates/groups' in line_lower or 'gangster' in line_lower:
            current_section = 'associates'
            continue
        elif 'part-iv' in line_lower or 'points for follow' in line_lower:
            current_section = 'narrative'
            continue
        elif 'part-1' in line_lower or 'personal details' in line_lower:
            current_section = 'personal'
            continue
        elif 'friends/associates' in line_lower:
            current_section = 'friends'
            continue

        # In associates section — each line is a name
        if current_section == 'associates':
            if line_stripped and not line_stripped.startswith('Anti') and not line_stripped.startswith('Places'):
                _add_associate(line_stripped, 'associate')
            continue

        if current_section == 'friends':
            if line_stripped:
                _add_associate(line_stripped, 'friend')
            continue

        # Parse "Field: Value" or "Field | Value" patterns
        # Handle pipe-separated table rows
        if '|' in line_stripped:
            parts = [p.strip() for p in line_stripped.split('|')]
            if len(parts) >= 2:
                field_raw = parts[0]
                # Take the last non-empty value (tables repeat values across columns)
                value_raw = next((p for p in reversed(parts[1:]) if p.strip()), '')
                if field_raw and value_raw:
                    field_lower = field_raw.lower().strip()
                    # Skip unwanted fields
                    if any(s in field_lower for s in skip_fields):
                        continue
                    val = _clean(value_raw)
                    if val and val.lower() not in ('nil', 'n/a', 'na', 'unknown', '-', 'none', ''):
                        field_name = _re.sub(r'\s+', ' ', field_raw).title()
                        # Check if it's a family field
                        for fk, rel in family_fields.items():
                            if fk in field_lower:
                                _add_family(val, rel)
                                break
                        # Check if it's an associate field
                        for ak, rel in associate_fields.items():
                            if ak in field_lower:
                                # Multiple names separated by newlines
                                for name_line in val.split('\n'):
                                    name_line = name_line.strip()
                                    if name_line:
                                        _add_associate(name_line, rel)
                                break
                        else:
                            # Regular personal field
                            if field_name not in personal:
                                personal[field_name] = val
            continue

        # Parse "Field: Value" lines
        m = _re.match(r'^([A-Za-z][A-Za-z\s/\(\)\.]{1,50}?)\s*[:\-]\s*(.+)$', line_stripped, _re.IGNORECASE)
        if m:
            field_raw = m.group(1).strip()
            value_raw = m.group(2).strip()
            field_lower = field_raw.lower().strip()

            if any(s in field_lower for s in skip_fields):
                continue

            val = _clean(value_raw)
            if not val or val.lower() in ('nil', 'n/a', 'na', 'unknown', '-', 'none', ''):
                continue

            field_name = _re.sub(r'\s+', ' ', field_raw).title()

            # Family fields
            for fk, rel in family_fields.items():
                if fk in field_lower:
                    _add_family(val, rel)
                    break
            # Associate fields
            for ak, rel in associate_fields.items():
                if ak in field_lower:
                    for name_line in val.split('\n'):
                        if name_line.strip():
                            _add_associate(name_line.strip(), rel)
                    break
            else:
                if field_name not in personal:
                    personal[field_name] = val

    # ── Step 2: Extract from narrative (PART-IV) ──────────────────────────
    # Find names mentioned with S/O, R/O patterns in narrative
    narrative_start = text.find('PART – IV')
    if narrative_start == -1:
        narrative_start = text.find('PART-IV')
    if narrative_start == -1:
        narrative_start = text.find('Points for follow')

    if narrative_start != -1:
        narrative = text[narrative_start:]
        # Extract S/O patterns (family)
        for m in _re.finditer(r'([A-Z][a-zA-Z\s@]{2,30})\s+S/O\s+([A-Z][a-zA-Z\s]{2,30})', narrative):
            person = _take_name(m.group(1))
            parent = _take_name(m.group(2))
            if person and parent and person.lower() not in seen:
                _add_associate(person, 'associate')
            if parent and parent.lower() not in seen:
                _add_family(parent, 'father')

        # Extract R/O patterns (associates with location)
        for m in _re.finditer(r'([A-Z][a-zA-Z\s@]{2,30})\s+R/[Oo]\s+([A-Z][a-zA-Z\s,]{2,60})', narrative):
            person = _take_name(m.group(1))
            if person and person.lower() not in seen:
                location = m.group(2).strip()[:50]
                _add_associate(person, 'associate', f'R/O {location}')

    # ── Step 3: Get main person name ──────────────────────────────────────
    main_person = (
        personal.get('Name', '') or
        personal.get('Accused', '') or
        personal.get('Victim', '') or
        personal.get('Complainant', '')
    )

    logger.info(f"[ProfileGraph] Extracted: {len(personal)} personal, {len(family)} family, {len(associates)} associates")
    logger.info(f"[ProfileGraph] Personal keys: {list(personal.keys())[:10]}")
    logger.info(f"[ProfileGraph] Family: {[f['name'] for f in family]}")
    logger.info(f"[ProfileGraph] Associates: {[a['name'] for a in associates[:5]]}")

    return {
        'personal': personal,
        'family': family,
        'associates': associates,
        'main_person': main_person,
    }

@router.get("/sources/{source_id}", response_model=SourceResponse)
async def get_source(source_id: str, include_text: bool = True):
    """Get a specific source by ID. Pass include_text=false to skip full_text for faster loads."""
    try:
        source = await Source.get(source_id)
        if not source:
            raise HTTPException(status_code=404, detail="Source not found")

        # Get status information if command exists
        status = None
        processing_info = None
        if source.command:
            try:
                status = await source.get_status()
                processing_info = await source.get_processing_progress()
            except Exception as e:
                logger.warning(f"Failed to get status for source {source_id}: {e}")
                status = "unknown"

        embedded_chunks = await source.get_embedded_chunks()

        # Get associated notebooks
        notebooks_query = await repo_query(
            "SELECT VALUE out FROM reference WHERE in = $source_id",
            {"source_id": ensure_record_id(source.id or source_id)},
        )
        notebook_ids = (
            [str(nb_id) for nb_id in notebooks_query] if notebooks_query else []
        )

        return SourceResponse(
            id=source.id or "",
            title=source.title,
            topics=source.topics or [],
            asset=AssetModel(
                file_path=source.asset.file_path if source.asset else None,
                url=source.asset.url if source.asset else None,
            )
            if source.asset
            else None,
            full_text=source.full_text if include_text else None,
            embedded=embedded_chunks > 0,
            embedded_chunks=embedded_chunks,
            file_available=_is_source_file_available(source),
            created=str(source.created),
            updated=str(source.updated),
            # Status fields
            command_id=str(source.command) if source.command else None,
            status=status,
            processing_info=processing_info,
            # Notebook associations
            notebooks=notebook_ids,
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error fetching source {source_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error fetching source: {str(e)}")


@router.head("/sources/{source_id}/download")
async def check_source_file(source_id: str):
    """Check if a source has a downloadable file."""
    try:
        await _resolve_source_file(source_id)
        return Response(status_code=200)
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error checking file for source {source_id}: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to verify file")


@router.get("/sources/{source_id}/download")
async def download_source_file(source_id: str):
    """Download the original file associated with an uploaded source."""
    try:
        resolved_path, filename = await _resolve_source_file(source_id)
        return FileResponse(
            path=resolved_path,
            filename=filename,
            media_type="application/octet-stream",
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error downloading file for source {source_id}: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to download source file")


@router.get("/sources/{source_id}/status", response_model=SourceStatusResponse)
async def get_source_status(source_id: str):
    """Get processing status for a source."""
    try:
        # First, verify source exists
        source = await Source.get(source_id)
        if not source:
            raise HTTPException(status_code=404, detail="Source not found")

        # Check if this is a legacy source (no command)
        if not source.command:
            return SourceStatusResponse(
                status=None,
                message="Legacy source (completed before async processing)",
                processing_info=None,
                command_id=None,
            )

        # Get command status and processing info
        try:
            status = await source.get_status()
            processing_info = await source.get_processing_progress()

            # Generate descriptive message based on status
            if status == "completed":
                message = "Source processing completed successfully"
            elif status == "failed":
                message = "Source processing failed"
            elif status == "running":
                message = "Source processing in progress"
            elif status == "queued":
                message = "Source processing queued"
            elif status == "unknown":
                message = "Source processing status unknown"
            else:
                message = f"Source processing status: {status}"

            return SourceStatusResponse(
                status=status,
                message=message,
                processing_info=processing_info,
                command_id=str(source.command) if source.command else None,
            )

        except Exception as e:
            logger.warning(f"Failed to get status for source {source_id}: {e}")
            return SourceStatusResponse(
                status="unknown",
                message="Failed to retrieve processing status",
                processing_info=None,
                command_id=str(source.command) if source.command else None,
            )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error fetching status for source {source_id}: {str(e)}")
        raise HTTPException(
            status_code=500, detail=f"Error fetching source status: {str(e)}"
        )


@router.put("/sources/{source_id}", response_model=SourceResponse)
async def update_source(source_id: str, source_update: SourceUpdate):
    """Update a source."""
    try:
        source = await Source.get(source_id)
        if not source:
            raise HTTPException(status_code=404, detail="Source not found")

        # Update only provided fields
        if source_update.title is not None:
            source.title = source_update.title
        if source_update.topics is not None:
            source.topics = source_update.topics

        await source.save()

        embedded_chunks = await source.get_embedded_chunks()
        return SourceResponse(
            id=source.id or "",
            title=source.title,
            topics=source.topics or [],
            asset=AssetModel(
                file_path=source.asset.file_path if source.asset else None,
                url=source.asset.url if source.asset else None,
            )
            if source.asset
            else None,
            full_text=source.full_text,
            embedded=embedded_chunks > 0,
            embedded_chunks=embedded_chunks,
            created=str(source.created),
            updated=str(source.updated),
        )
    except HTTPException:
        raise
    except InvalidInputError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Error updating source {source_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error updating source: {str(e)}")


@router.post("/sources/{source_id}/retry", response_model=SourceResponse)
async def retry_source_processing(source_id: str):
    """Retry processing for a failed or stuck source."""
    try:
        # First, verify source exists
        source = await Source.get(source_id)
        if not source:
            raise HTTPException(status_code=404, detail="Source not found")

        # Check if source already has a running command
        if source.command:
            try:
                status = await source.get_status()
                if status in ["running", "queued"]:
                    raise HTTPException(
                        status_code=400,
                        detail="Source is already processing. Cannot retry while processing is active.",
                    )
            except Exception as e:
                logger.warning(
                    f"Failed to check current status for source {source_id}: {e}"
                )
                # Continue with retry if we can't check status

        # Get notebooks that this source belongs to
        query = "SELECT notebook FROM reference WHERE source = $source_id"
        references = await repo_query(query, {"source_id": source_id})
        notebook_ids = [str(ref["notebook"]) for ref in references]

        if not notebook_ids:
            raise HTTPException(
                status_code=400, detail="Source is not associated with any notebooks"
            )

        # Prepare content_state based on source asset
        content_state = {}
        if source.asset:
            if source.asset.file_path:
                content_state = {
                    "file_path": source.asset.file_path,
                    "delete_source": False,  # Don't delete on retry
                }
            elif source.asset.url:
                content_state = {"url": source.asset.url}
            else:
                raise HTTPException(
                    status_code=400, detail="Source asset has no file_path or url"
                )
        else:
            # Check if it's a text source by trying to get full_text
            if source.full_text:
                content_state = {"content": source.full_text}
            else:
                raise HTTPException(
                    status_code=400, detail="Cannot determine source content for retry"
                )

        try:
            # Import command modules to ensure they're registered
            import commands.source_commands  # noqa: F401

            # Submit new command for background processing
            command_input = SourceProcessingInput(
                source_id=str(source.id),
                content_state=content_state,
                notebook_ids=notebook_ids,
                transformations=[],  # Use default transformations on retry
                embed=True,  # Always embed on retry
            )

            command_id = await CommandService.submit_command_job(
                "open_notebook",  # app name
                "process_source",  # command name
                command_input.model_dump(),
            )

            logger.info(
                f"Submitted retry processing command: {command_id} for source {source_id}"
            )

            # Update source with new command ID
            source.command = ensure_record_id(f"command:{command_id}")
            await source.save()

            # Get current embedded chunks count
            embedded_chunks = await source.get_embedded_chunks()

            # Return updated source response
            return SourceResponse(
                id=source.id or "",
                title=source.title,
                topics=source.topics or [],
                asset=AssetModel(
                    file_path=source.asset.file_path if source.asset else None,
                    url=source.asset.url if source.asset else None,
                )
                if source.asset
                else None,
                full_text=source.full_text,
                embedded=embedded_chunks > 0,
                embedded_chunks=embedded_chunks,
                created=str(source.created),
                updated=str(source.updated),
                command_id=command_id,
                status="queued",
                processing_info={"retry": True, "queued": True},
            )

        except Exception as e:
            logger.error(
                f"Failed to submit retry processing command for source {source_id}: {e}"
            )
            raise HTTPException(
                status_code=500, detail=f"Failed to queue retry processing: {str(e)}"
            )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error retrying source processing for {source_id}: {str(e)}")
        raise HTTPException(
            status_code=500, detail=f"Error retrying source processing: {str(e)}"
        )


@router.delete("/sources/{source_id}")
async def delete_source(source_id: str):
    """Delete a source."""
    try:
        source = await Source.get(source_id)
        if not source:
            raise HTTPException(status_code=404, detail="Source not found")

        await source.delete()

        return {"message": "Source deleted successfully"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting source {source_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error deleting source: {str(e)}")


@router.get("/sources/{source_id}/insights", response_model=List[SourceInsightResponse])
async def get_source_insights(source_id: str):
    """Get all insights for a specific source."""
    try:
        source = await Source.get(source_id)
        if not source:
            raise HTTPException(status_code=404, detail="Source not found")

        insights = await source.get_insights()
        return [
            SourceInsightResponse(
                id=insight.id or "",
                source_id=source_id,
                insight_type=insight.insight_type,
                content=insight.content,
                created=str(insight.created),
                updated=str(insight.updated),
            )
            for insight in insights
        ]
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error fetching insights for source {source_id}: {str(e)}")
        raise HTTPException(
            status_code=500, detail=f"Error fetching insights: {str(e)}"
        )


@router.delete("/sources/{source_id}/insights/mindmap")
async def delete_mindmap_insights(source_id: str):
    """Delete all Mind Map insights for a source so they can be regenerated fresh."""
    try:
        source = await Source.get(source_id)
        if not source:
            raise HTTPException(status_code=404, detail="Source not found")

        insights = await source.get_insights()
        deleted = 0
        for insight in insights:
            if insight.insight_type and 'mind' in insight.insight_type.lower():
                await insight.delete()
                deleted += 1

        return {"message": f"Deleted {deleted} Mind Map insight(s)", "deleted": deleted}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting mindmap insights for source {source_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post(
    "/sources/{source_id}/insights",
    response_model=InsightCreationResponse,
    status_code=202,
)
async def create_source_insight(source_id: str, request: CreateSourceInsightRequest):
    """
    Start insight generation for a source by running a transformation.

    This endpoint returns immediately with a 202 Accepted status.
    The transformation runs asynchronously in the background via the job queue.
    Poll GET /sources/{source_id}/insights to see when the insight is ready.
    """
    try:
        # Validate source exists
        source = await Source.get(source_id)
        if not source:
            raise HTTPException(status_code=404, detail="Source not found")

        # Validate transformation exists
        transformation = await Transformation.get(request.transformation_id)
        if not transformation:
            raise HTTPException(status_code=404, detail="Transformation not found")

        # Submit transformation as background job (fire-and-forget)
        command_id = submit_command(
            "open_notebook",
            "run_transformation",
            {
                "source_id": source_id,
                "transformation_id": request.transformation_id,
            },
        )
        logger.info(
            f"Submitted run_transformation command {command_id} for source {source_id}"
        )

        # Return immediately with command_id for status tracking
        return InsightCreationResponse(
            status="pending",
            message="Insight generation started",
            source_id=source_id,
            transformation_id=request.transformation_id,
            command_id=str(command_id),
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error starting insight generation for source {source_id}: {e}")
        raise HTTPException(
            status_code=500, detail=f"Error starting insight generation: {str(e)}"
        )

