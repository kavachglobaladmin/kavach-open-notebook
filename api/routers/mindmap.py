# import asyncio
# import base64
# import os
# from typing import Any, Dict, List, Optional
# from urllib.parse import unquote
# import json
# import re

# import numpy as np
# from fastapi import APIRouter, HTTPException
# from loguru import logger
# from pydantic import BaseModel

# from open_notebook.domain.notebook import Source

# router = APIRouter()

# KAFKA_BOOTSTRAP_SERVERS = os.environ.get("KAFKA_BOOTSTRAP_SERVERS", "kafka:9093")


# def _decode_source_id(source_id: str) -> str:
#     """URL-decode source_id — FastAPI :path doesn't auto-decode %3A → : """
#     return unquote(source_id)

# # Module-level orchestrator instance (initialized on first use)
# _orchestrator: Optional[Any] = None


# def _build_orchestrator():
#     """Build and return a KafkaMindMapOrchestrator with a fully wired MindMapPipeline."""
#     from langchain_ollama import ChatOllama
#     from open_notebook.graphs.mind_map import (
#         EasyOCRService,
#         IntelligenceLLMService,
#         KafkaMindMapOrchestrator,
#         MindMapPipeline,
#         TextProcessor,
#     )

#     # Read env vars fresh at build time — never at module load — so the correct
#     # runtime value is always used even after container restarts or hot-reloads.
#     ollama_url = os.environ.get("OLLAMA_BASE_URL", "http://host.docker.internal:11434")
#     kafka_servers = os.environ.get("KAFKA_BOOTSTRAP_SERVERS", "kafka:9093")

#     logger.info(f"Building MindMapPipeline — Ollama: {ollama_url}, Kafka: {kafka_servers}")
#     llm = ChatOllama(model="qwen3", temperature=0.2, base_url=ollama_url)
#     ocr_service = EasyOCRService()
#     text_processor = TextProcessor()
#     llm_service = IntelligenceLLMService(llm)
#     pipeline = MindMapPipeline(
#         ocr_service=ocr_service,
#         processor=text_processor,
#         llm_service=llm_service,
#     )
#     orchestrator = KafkaMindMapOrchestrator(
#         pipeline=pipeline,
#         bootstrap_servers=kafka_servers,
#     )
#     logger.info("KafkaMindMapOrchestrator ready")
#     return orchestrator


# def get_orchestrator():
#     global _orchestrator
#     if _orchestrator is None:
#         _orchestrator = _build_orchestrator()
#     return _orchestrator


# class MindMapRequest(BaseModel):
#     model_name: str = "qwen3"
#     temperature: float = 0.2


# class MindMapResponse(BaseModel):
#     mind_map: Dict[str, Any]
#     source_id: str


# @router.post("/sources/{source_id}/mindmap", response_model=MindMapResponse)
# async def generate_mind_map(source_id: str, request: MindMapRequest):
#     """Generate a mind map from a source's content using KafkaMindMapOrchestrator."""
#     try:
#         source_id = _decode_source_id(source_id)
#         try:
#             source = await Source.get(source_id)
#         except Exception:
#             raise HTTPException(status_code=404, detail=f"Source not found: {source_id}")

#         if not source.full_text or not source.full_text.strip():
#             raise HTTPException(
#                 status_code=400,
#                 detail="Source has no text content to generate a mind map from",
#             )

#         orchestrator = get_orchestrator()

#         # Publish job to Kafka (non-blocking, best-effort)
#         asyncio.create_task(_safe_produce(orchestrator, source_id))

#         # Generate directly via the pipeline (no timeout)
#         logger.info(f"Generating mind map for source_id={source_id}")
#         mind_map = await orchestrator.pipeline.generate_from_source_id(source_id)
        
#         # Validate the response is a proper dict
#         if not isinstance(mind_map, dict):
#             logger.error(f"Mind map returned non-dict type: {type(mind_map)}")
#             raise ValueError("Mind map generation returned invalid type")
        
#         if not mind_map.get("label"):
#             logger.error("Mind map missing required 'label' field")
#             mind_map = {
#                 "label": "Generated Mind Map",
#                 "children": [{"label": "Unable to generate structured mind map"}]
#             }
        
#         logger.success(f"Mind map generation completed for source_id={source_id}")
#         return MindMapResponse(mind_map=mind_map, source_id=source_id)

#     except HTTPException:
#         raise
#     except ValueError as e:
#         logger.error(f"Mind map validation error: {e}")
#         raise HTTPException(status_code=422, detail=str(e))
#     except Exception as e:
#         logger.error(f"Mind map generation failed for source {source_id}: {type(e).__name__}: {e}")
#         import traceback
#         logger.debug(f"Traceback: {traceback.format_exc()}")
#         raise HTTPException(status_code=500, detail=f"Mind map generation failed: {str(e)}")


# async def _safe_produce(orchestrator, source_id: str):
#     """Fire-and-forget Kafka job publish — errors are logged, never raised."""
#     try:
#         await orchestrator.produce_jobs([source_id])
#     except Exception as e:
#         logger.warning(f"Kafka produce skipped for {source_id}: {e}")


# async def start_kafka_consumer():
#     """Start the KafkaMindMapOrchestrator consumer as a background task."""
#     try:
#         orchestrator = get_orchestrator()
#         logger.info("Starting KafkaMindMapOrchestrator consumer...")
#         await orchestrator.start_consumer()
#     except Exception as e:
#         logger.warning(f"Kafka consumer could not start (Kafka may be unavailable): {e}")


# class SourceImagesResponse(BaseModel):
#     images: List[str]  # base64-encoded PNG strings
#     source_id: str
#     count: int


# class NodeSummaryRequest(BaseModel):
#     node_name: str
#     root_subject: str  # the root node label (person/topic name)


# class NodeSummaryResponse(BaseModel):
#     summary: str
#     node_name: str
#     root_subject: str


# @router.post("/sources/{source_id}/node-summary", response_model=NodeSummaryResponse)
# async def get_node_summary(source_id: str, request: NodeSummaryRequest):
#     """Generate a detailed summary for a specific mind map node using the source content."""
#     try:
#         logger.info(f"Node summary request: source_id={source_id!r}, node={request.node_name!r}")
#         source_id = _decode_source_id(source_id)
#         logger.info(f"Decoded source_id={source_id!r}")
#         try:
#             source = await Source.get(source_id)
#         except Exception:
#             raise HTTPException(status_code=404, detail=f"Source not found: {source_id}")

#         if not source:
#             raise HTTPException(status_code=404, detail="Source not found")

#         if not source.full_text or not source.full_text.strip():
#             raise HTTPException(status_code=400, detail="Source has no text content")

#         ollama_url = os.environ.get("OLLAMA_BASE_URL", "http://host.docker.internal:11434")

#         from langchain_ollama import ChatOllama
#         from langchain_core.messages import HumanMessage, SystemMessage

#         llm = ChatOllama(model="qwen3", temperature=0.3, base_url=ollama_url)

#         # Truncate source text to avoid context overflow (~12k chars)
#         context_text = source.full_text[:12000]

#         system_prompt = (
#             "You are an expert analyst. Given source document content, provide a detailed, "
#             "well-structured summary about a specific topic as it relates to the main subject. "
#             "Be thorough, cite specific facts from the source, and organize your response clearly. "
#             "Do not add information not present in the source. "
#             "STRICT FORMATTING RULES: "
#             "- Do NOT use any markdown headers (no #, ##, ### etc). "
#             "- Do NOT use --- horizontal rules. "
#             "- Use **bold** only for section titles and key terms. "
#             "- Use plain numbered lists or bullet points for structure. "
#             "- Write in clean plain text paragraphs."
#         )

#         user_prompt = (
#             f"Discuss what these sources say about '{request.node_name}', "
#             f"in the larger context of '{request.root_subject}'.\n\n"
#             f"Source content:\n{context_text}"
#         )

#         def _run_llm():
#             messages = [
#                 SystemMessage(content=system_prompt),
#                 HumanMessage(content=user_prompt),
#             ]
#             response = llm.invoke(messages)
#             return response.content

#         summary = await asyncio.to_thread(_run_llm)

#         # Strip <think>...</think> tags if model outputs chain-of-thought
#         import re
#         summary = re.sub(r'<think>.*?</think>', '', summary, flags=re.DOTALL).strip()
#         # Strip markdown headers (###, ##, #) — replace with bold text instead
#         summary = re.sub(r'^#{1,6}\s+(.+)$', r'**\1**', summary, flags=re.MULTILINE)
#         # Strip horizontal rules
#         summary = re.sub(r'^-{3,}$', '', summary, flags=re.MULTILINE)
#         # Collapse multiple blank lines
#         summary = re.sub(r'\n{3,}', '\n\n', summary).strip()

#         logger.info(f"Node summary generated for '{request.node_name}' in source {source_id}")
#         return NodeSummaryResponse(
#             summary=summary,
#             node_name=request.node_name,
#             root_subject=request.root_subject,
#         )

#     except HTTPException:
#         raise
#     except Exception as e:
#         logger.error(f"Node summary failed for source {source_id}: {e}")
#         raise HTTPException(status_code=500, detail=f"Node summary failed: {str(e)}")


# class SourceSummaryResponse(BaseModel):
#     summary: str
#     source_id: str


# @router.post("/sources/{source_id}/summary", response_model=SourceSummaryResponse)
# async def get_source_summary(source_id: str):
#     """Generate a full summary of the source document using Ollama qwen3 via SummaryPipeline."""
#     try:
#         source_id = _decode_source_id(source_id)
#         logger.info(f"Source summary request: source_id={source_id!r}")

#         from langchain_ollama import ChatOllama
#         from open_notebook.graphs.summary import SummaryPipeline, SummaryTextProcessor, SummaryLLMService

#         ollama_url = os.environ.get("OLLAMA_BASE_URL", "http://host.docker.internal:11434")
#         llm = ChatOllama(model="qwen3", temperature=0.3, base_url=ollama_url)

#         pipeline = SummaryPipeline(
#             processor=SummaryTextProcessor(),
#             llm_service=SummaryLLMService(llm),
#         )

#         result = await pipeline.generate_from_source_id(source_id)

#         logger.info(f"Source summary generated for source {source_id}")
#         return SourceSummaryResponse(summary=result["summary"], source_id=source_id)

#     except HTTPException:
#         raise
#     except Exception as e:
#         logger.error(f"Source summary failed for source {source_id}: {e}")
#         raise HTTPException(status_code=500, detail=f"Source summary failed: {str(e)}")


# def _extract_images_from_pdf(file_path: str) -> List[str]:
#     """
#     Extract images from a PDF using two strategies:
#     1. Embedded XObject images (photos embedded in the PDF stream)
#     2. Page rendering — render each page as a high-res image (catches scanned PDFs,
#        pages with photos drawn directly, charts, etc.)
#     Returns a deduplicated list of base64-encoded PNG strings.
#     """
#     import io
#     import fitz  # PyMuPDF
#     from PIL import Image

#     images_b64: List[str] = []
#     seen_sizes: set = set()  # deduplicate by (width, height, first-bytes)

#     doc = fitz.open(file_path)

#     for page_index in range(len(doc)):
#         page = doc[page_index]

#         # ── Strategy 1: extract embedded XObject images ──────────────────────
#         embedded = page.get_images(full=True)
#         page_has_embedded = False

#         for img_info in embedded:
#             xref = img_info[0]
#             try:
#                 base_image = doc.extract_image(xref)
#                 img_bytes = base_image["image"]
#                 img = Image.open(io.BytesIO(img_bytes))
#                 # Skip tiny icons / artifacts
#                 if img.width < 80 or img.height < 80:
#                     continue
#                 buf = io.BytesIO()
#                 img.convert("RGB").save(buf, format="PNG")
#                 raw = buf.getvalue()
#                 key = (img.width, img.height, raw[:64])
#                 if key in seen_sizes:
#                     continue
#                 seen_sizes.add(key)
#                 images_b64.append(base64.b64encode(raw).decode("utf-8"))
#                 page_has_embedded = True
#             except Exception as e:
#                 logger.warning(f"Skipping embedded image xref={xref} on page {page_index}: {e}")

#         # ── Strategy 2: render the full page if it has no embedded images ────
#         # This catches scanned PDFs and pages where photos are drawn as page content
#         if not page_has_embedded:
#             try:
#                 # 150 DPI is a good balance of quality vs size
#                 mat = fitz.Matrix(150 / 72, 150 / 72)
#                 pix = page.get_pixmap(matrix=mat, alpha=False)
#                 img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

#                 # Skip pages that are mostly white (text-only pages)
#                 arr = np.array(img)
#                 white_ratio = (arr > 240).all(axis=2).mean()
#                 if white_ratio > 0.92:
#                     # Mostly blank/text page — skip
#                     continue

#                 buf = io.BytesIO()
#                 img.save(buf, format="PNG")
#                 raw = buf.getvalue()
#                 key = (img.width, img.height, raw[:64])
#                 if key not in seen_sizes:
#                     seen_sizes.add(key)
#                     images_b64.append(base64.b64encode(raw).decode("utf-8"))
#             except Exception as e:
#                 logger.warning(f"Page render failed for page {page_index}: {e}")

#     doc.close()
#     return images_b64


# def _extract_images_from_docx(file_path: str) -> List[str]:
#     """Extract embedded images from a .docx file using python-docx."""
#     import io
#     from PIL import Image
#     from docx import Document

#     images_b64: List[str] = []
#     seen: set = set()

#     doc = Document(file_path)
#     for rel in doc.part.rels.values():
#         if "image" in rel.reltype:
#             try:
#                 img_bytes = rel.target_part.blob
#                 img = Image.open(io.BytesIO(img_bytes))
#                 if img.width < 80 or img.height < 80:
#                     continue
#                 buf = io.BytesIO()
#                 img.convert("RGB").save(buf, format="PNG")
#                 raw = buf.getvalue()
#                 key = raw[:128]
#                 if key in seen:
#                     continue
#                 seen.add(key)
#                 images_b64.append(base64.b64encode(raw).decode("utf-8"))
#             except Exception as e:
#                 logger.warning(f"Skipping docx image: {e}")

#     return images_b64


# @router.get("/sources/{source_id}/images", response_model=SourceImagesResponse)
# async def get_source_images(source_id: str):
#     """Extract images from a source's PDF/file and return them as base64 PNGs."""
#     try:
#         source_id = _decode_source_id(source_id)
#         try:
#             source = await Source.get(source_id)
#         except Exception:
#             raise HTTPException(status_code=404, detail=f"Source not found: {source_id}")

#         file_path = source.asset.file_path if source.asset else None
#         if not file_path or not os.path.exists(file_path):
#             # No file — return empty list gracefully
#             return SourceImagesResponse(images=[], source_id=source_id, count=0)

#         ext = os.path.splitext(file_path)[1].lower()
#         images_b64: List[str] = []

#         if ext == ".pdf":
#             # Run in thread pool — CPU-bound PDF rendering
#             images_b64 = await asyncio.to_thread(_extract_images_from_pdf, file_path)
#         elif ext in (".docx", ".doc"):
#             images_b64 = await asyncio.to_thread(_extract_images_from_docx, file_path)
#         else:
#             # For image files themselves, return the file directly
#             try:
#                 import io
#                 from PIL import Image
#                 img = Image.open(file_path)
#                 buf = io.BytesIO()
#                 img.convert("RGB").save(buf, format="PNG")
#                 images_b64.append(base64.b64encode(buf.getvalue()).decode("utf-8"))
#             except Exception as e:
#                 logger.warning(f"Could not read file as image for {source_id}: {e}")

#         logger.info(f"Extracted {len(images_b64)} images from source {source_id}")
#         return SourceImagesResponse(images=images_b64, source_id=source_id, count=len(images_b64))

#     except HTTPException:
#         raise
#     except Exception as e:
#         logger.error(f"Image extraction failed for source {source_id}: {e}")
#         raise HTTPException(status_code=500, detail=f"Image extraction failed: {str(e)}")




# import asyncio
# import base64
# import os
# import json
# import re
# from typing import Any, Dict, List, Optional
# from urllib.parse import unquote

# import numpy as np
# from fastapi import APIRouter, HTTPException
# from loguru import logger
# from pydantic import BaseModel

# from open_notebook.domain.notebook import Source

# router = APIRouter()

# KAFKA_BOOTSTRAP_SERVERS = os.environ.get("KAFKA_BOOTSTRAP_SERVERS", "kafka:9093")


# # ---------------------------------------------------------------------------
# # Helpers
# # ---------------------------------------------------------------------------

# def _decode_source_id(source_id: str) -> str:
#     """URL-decode source_id — FastAPI :path doesn't auto-decode %3A → : """
#     return unquote(source_id)


# def _parse_llm_json(raw: str) -> dict:
#     """
#     Robustly parse JSON from an LLM response that may contain:
#       - <think>...</think> chain-of-thought blocks
#       - Markdown code fences:  ```json { ... } ```  or  ``` { ... } ```
#       - Python function syntax: json(...), json.loads(...), etc.
#       - Leading/trailing prose around the JSON object
#       - Other formatting noise

#     Returns a clean dict, raising ValueError if no valid JSON found.
#     """
#     logger.debug(f"[_parse_llm_json] Raw input (first 300 chars): {raw[:300]}")

#     original_raw = raw

#     # 1. Strip <think>...</think> blocks
#     cleaned = re.sub(r"<think>[\s\S]*?</think>", "", raw, flags=re.DOTALL).strip()

#     # 2. Remove common function call wrappers
#     cleaned = re.sub(r'^\s*(?:json\.loads\s*\(|json\s*\(|python\s*\()', '', cleaned, flags=re.IGNORECASE)
#     cleaned = re.sub(r'\)\s*$', '', cleaned).strip()

#     # 3. Try markdown code fence
#     fence_match = re.search(r"```(?:json)?\s*([\s\S]*?)\s*```", cleaned, flags=re.DOTALL)
#     if fence_match:
#         fence_content = fence_match.group(1).strip()
#         try:
#             result = json.loads(fence_content)
#             if isinstance(result, dict) and result.get("label"):
#                 logger.debug("[_parse_llm_json] ✓ Parsed from markdown fence")
#                 return result
#         except json.JSONDecodeError:
#             pass

#     # 4. Find ALL { } candidates (bracket-balanced search)
#     candidates = []
#     depth = 0
#     start = -1
#     in_string = False
#     escape_next = False

#     for i, char in enumerate(cleaned):
#         if escape_next:
#             escape_next = False
#             continue
#         if char == '\\':
#             escape_next = True
#             continue
#         if char == '"' and depth >= 0:
#             in_string = not in_string
#             continue
#         if in_string:
#             continue
#         if char == '{':
#             if depth == 0:
#                 start = i
#             depth += 1
#         elif char == '}':
#             depth -= 1
#             if depth == 0 and start != -1:
#                 candidate = cleaned[start: i + 1]
#                 try:
#                     parsed = json.loads(candidate)
#                     if isinstance(parsed, dict) and parsed.get("label"):
#                         candidates.append(parsed)
#                 except json.JSONDecodeError:
#                     pass
#                 start = -1

#     if candidates:
#         best = sorted(candidates, key=lambda x: len(json.dumps(x)), reverse=True)[0]
#         logger.debug(f"[_parse_llm_json] ✓ Parsed from {len(candidates)} candidates")
#         return best

#     # 5. Try parsing the whole cleaned string
#     try:
#         result = json.loads(cleaned)
#         if isinstance(result, dict) and result.get("label"):
#             return result
#     except json.JSONDecodeError:
#         pass

#     logger.error(f"[_parse_llm_json] FAILED to parse. Raw (500 chars):\n{original_raw[:500]}\n...")
#     raise ValueError(f"Could not extract valid JSON with 'label' key. Raw preview: {original_raw[:150]}...")


# def _build_fallback_mind_map(text: str, title: str = "Mind Map") -> dict:
#     """
#     Build a basic mind map from plain text when LLM returns prose instead of JSON.
#     Splits text into sentences and groups them into generic categories.
#     """
#     sentences = [s.strip() for s in re.split(r'(?<=[.!?])\s+', text) if len(s.strip()) > 20]
#     chunk_size = max(3, len(sentences) // 5)
#     children = []
#     for i in range(0, len(sentences), chunk_size):
#         chunk = sentences[i:i + chunk_size]
#         children.append({
#             "label": f"Section {i // chunk_size + 1}",
#             "children": [{"label": s[:120]} for s in chunk]
#         })
#     return {
#         "label": title,
#         "children": children if children else [{"label": "No structured content found"}]
#     }


# # ---------------------------------------------------------------------------
# # Orchestrator (module-level singleton)
# # ---------------------------------------------------------------------------

# _orchestrator: Optional[Any] = None


# def _build_orchestrator():
#     """Build and return a KafkaMindMapOrchestrator with a fully wired MindMapPipeline."""
#     from langchain_ollama import ChatOllama
#     from open_notebook.graphs.mind_map import (
#         EasyOCRService,
#         IntelligenceLLMService,
#         KafkaMindMapOrchestrator,
#         MindMapPipeline,
#         TextProcessor,
#     )

#     ollama_url = os.environ.get("OLLAMA_BASE_URL", "http://host.docker.internal:11434")
#     kafka_servers = os.environ.get("KAFKA_BOOTSTRAP_SERVERS", "kafka:9093")

#     # llama3 works best with temperature=0 for strict JSON output
#     ollama_model = os.environ.get("OLLAMA_MODEL", os.environ.get("DEFAULT_CHAT_MODEL", "llama3:latest"))
#     print(f"Using Ollama model: {ollama_model}")
#     logger.info(f"Building MindMapPipeline — model: {ollama_model}, Ollama: {ollama_url}, Kafka: {kafka_servers}")

#     llm = ChatOllama(
#         model=ollama_model,
#         temperature=0,           # llama3: temperature=0 gives most consistent JSON
#         base_url=ollama_url,
#         num_ctx=8192,            # llama3 Large Ctx — use available context
#         repeat_penalty=1.1,      # reduce repetition in output
#     )
#     ocr_service = EasyOCRService()
#     text_processor = TextProcessor()
#     llm_service = IntelligenceLLMService(llm)
#     pipeline = MindMapPipeline(
#         ocr_service=ocr_service,
#         processor=text_processor,
#         llm_service=llm_service,
#     )
#     orchestrator = KafkaMindMapOrchestrator(
#         pipeline=pipeline,
#         bootstrap_servers=kafka_servers,
#     )
#     logger.info("KafkaMindMapOrchestrator ready")
#     return orchestrator


# def get_orchestrator():
#     global _orchestrator
#     if _orchestrator is None:
#         _orchestrator = _build_orchestrator()
#     return _orchestrator


# # ---------------------------------------------------------------------------
# # Models
# # ---------------------------------------------------------------------------

# class MindMapRequest(BaseModel):
#     model_name: str = "llama3:latest"
#     temperature: float = 0.0


# class MindMapResponse(BaseModel):
#     mind_map: Dict[str, Any]
#     source_id: str


# # ---------------------------------------------------------------------------
# # Mind Map endpoint
# # ---------------------------------------------------------------------------

# def _build_mindmap_from_docx(file_path: str) -> Optional[Dict[str, Any]]:
#     """
#     Build a clean, structured mind map directly from DOCX tables and PART-IV narrative.
#     This bypasses the LLM/OCR pipeline entirely for IR documents.
#     Returns None if the file is not a valid IR docx.
#     """
#     from api.routers.sources import _extract_profile_from_docx, _extract_part_iv_from_docx

#     profile = _extract_profile_from_docx(file_path)
#     part4 = _extract_part_iv_from_docx(file_path)

#     personal: dict = profile.get('personal', {})
#     family: list = profile.get('family', [])
#     associates: list = profile.get('associates', [])
#     main_person: str = profile.get('main_person', '') or personal.get('Name', 'Unknown Subject')
#     part4_sections: dict = part4.get('sections', {})

#     if not personal and not part4_sections:
#         return None  # Not a structured IR doc — fall back to LLM pipeline

#     # ── Helper ────────────────────────────────────────────────────────────
#     def _leaf(text: str) -> Dict:
#         return {"label": re.sub(r'\s+', ' ', text).strip()}

#     def _section(label: str, children: list) -> Dict:
#         return {"label": label, "children": children}

#     children: list = []

#     # ── 1. Personal Profile ───────────────────────────────────────────────
#     personal_leaves = []

#     # Priority fields in display order
#     priority_fields = [
#         'Name', 'Code Name', 'Alias', 'Date Of Birth', 'Age', 'Sex', 'Nationality',
#         'Religion', 'Caste', 'Marital Status', 'Complexion', 'Height', 'Weight',
#         'Build', 'Hair', 'Eyes', 'Descriptive Roll', 'Mark Of Identification',
#         'Education', 'Occupation', 'Bad Habits', 'Present Address', 'Permanent Address',
#         'Facebook Id', 'Email Id', 'Mobile No',
#     ]
#     shown_keys = set()
#     for field in priority_fields:
#         for k, v in personal.items():
#             if k.lower().strip() == field.lower() and v and v not in ('---', '-', 'N/A', 'Nil'):
#                 personal_leaves.append(_leaf(f"{k}: {v}"))
#                 shown_keys.add(k.lower())
#                 break

#     # Remaining personal fields not in priority list
#     for k, v in personal.items():
#         if k.lower() not in shown_keys and v and v not in ('---', '-', 'N/A', 'Nil'):
#             personal_leaves.append(_leaf(f"{k}: {v}"))

#     if personal_leaves:
#         children.append(_section("Personal Profile", personal_leaves))

#     # ── 2. Family & Relationships ─────────────────────────────────────────
#     family_leaves = []
#     for member in family:
#         name = member.get('name', '')
#         relation = member.get('relation', '').title()
#         details = member.get('details', '')
#         if name:
#             label = f"{relation}: {name}"
#             if details:
#                 label += f" ({details})"
#             family_leaves.append(_leaf(label))

#     if family_leaves:
#         children.append(_section("Family & Relationships", family_leaves))

#     # ── 3. Criminal Career (from PART-IV sections) ────────────────────────
#     criminal_leaves = []
#     criminal_keys = [
#         'how he got involved in crime', 'how she got involved in crime',
#         'criminal career', 'criminal history', 'modus operandi',
#         'expertise in criminal act', 'occupation before joining crime',
#         'weapons details', 'weapon details', 'bad habits',
#     ]
#     for sec_key, sec_val in part4_sections.items():
#         sk = sec_key.lower()
#         if any(ck in sk for ck in criminal_keys):
#             # Summarize long text into bullet sentences
#             sentences = [s.strip() for s in re.split(r'(?<=[.!?])\s+', sec_val) if len(s.strip()) > 15]
#             if sentences:
#                 criminal_leaves.append(_section(sec_key, [_leaf(s) for s in sentences[:6]]))
#             elif sec_val.strip():
#                 criminal_leaves.append(_leaf(f"{sec_key}: {sec_val[:120]}"))

#     # Gang affiliations from associates
#     gang_leaves = []
#     for assoc in associates:
#         rel = assoc.get('relation', '').lower()
#         if any(g in rel for g in ['gang', 'gangster', 'associate', 'group']):
#             name = assoc.get('name', '')
#             details = assoc.get('details', '')
#             label = name
#             if details:
#                 label += f" — {details}"
#             gang_leaves.append(_leaf(label))

#     if gang_leaves:
#         criminal_leaves.append(_section("Gang Associates", gang_leaves))

#     if criminal_leaves:
#         children.append(_section("Criminal Career", criminal_leaves))

#     # ── 4. Legal History (FIR cases from personal table) ─────────────────
#     legal_leaves = []
#     # Extract FIR data from personal fields
#     fir_fields = {k: v for k, v in personal.items()
#                   if any(x in k.lower() for x in ['fir', 'case', 'arrest', 'police station', 'status of case', 'under section', 'u/s'])}
#     for k, v in fir_fields.items():
#         legal_leaves.append(_leaf(f"{k}: {v}"))

#     # Also from PART-IV sections
#     legal_keys = ['previous involvements', 'legal history', 'fir details', 'case details', 'arrest details']
#     for sec_key, sec_val in part4_sections.items():
#         sk = sec_key.lower()
#         if any(lk in sk for lk in legal_keys):
#             sentences = [s.strip() for s in re.split(r'(?<=[.!?])\s+', sec_val) if len(s.strip()) > 15]
#             if sentences:
#                 legal_leaves.append(_section(sec_key, [_leaf(s) for s in sentences[:8]]))
#             elif sec_val.strip():
#                 legal_leaves.append(_leaf(f"{sec_key}: {sec_val[:150]}"))

#     if legal_leaves:
#         children.append(_section("Legal History", legal_leaves))

#     # ── 5. Movements & Hideouts ───────────────────────────────────────────
#     movement_leaves = []
#     movement_keys = ['movements', 'hideout', 'places', 'location', 'hide out', 'travel', 'stayed', 'residence']
#     for sec_key, sec_val in part4_sections.items():
#         sk = sec_key.lower()
#         if any(mk in sk for mk in movement_keys):
#             sentences = [s.strip() for s in re.split(r'(?<=[.!?])\s+', sec_val) if len(s.strip()) > 15]
#             if sentences:
#                 movement_leaves.append(_section(sec_key, [_leaf(s) for s in sentences[:6]]))
#             elif sec_val.strip():
#                 movement_leaves.append(_leaf(f"{sec_key}: {sec_val[:150]}"))

#     # Address fields
#     for k, v in personal.items():
#         if 'address' in k.lower() and v and v not in ('---', '-', 'N/A', 'Nil'):
#             movement_leaves.append(_leaf(f"{k}: {v}"))

#     if movement_leaves:
#         children.append(_section("Movements & Hideouts", movement_leaves))

#     # ── 6. Any remaining PART-IV sections ────────────────────────────────
#     handled_keys = set()
#     for sec_key in part4_sections:
#         sk = sec_key.lower()
#         already = any(
#             any(ck in sk for ck in criminal_keys + legal_keys + movement_keys)
#             for _ in [1]
#         )
#         if not already:
#             sec_val = part4_sections[sec_key]
#             sentences = [s.strip() for s in re.split(r'(?<=[.!?])\s+', sec_val) if len(s.strip()) > 15]
#             if sentences:
#                 children.append(_section(sec_key, [_leaf(s) for s in sentences[:5]]))
#             elif sec_val.strip():
#                 children.append(_leaf(f"{sec_key}: {sec_val[:120]}"))

#     if not children:
#         return None

#     return {
#         "label": main_person,
#         "children": children,
#     }


# @router.post("/sources/{source_id}/mindmap", response_model=MindMapResponse)
# async def generate_mind_map(source_id: str, request: MindMapRequest):
#     """Generate a mind map from a source's content.
    
#     Priority:
#     1. DOCX structured extraction (IR documents) — fast, accurate, no LLM needed
#     2. LLM pipeline fallback for non-DOCX or unstructured sources
#     """
#     try:
#         source_id = _decode_source_id(source_id)

#         try:
#             source = await Source.get(source_id)
#         except Exception:
#             raise HTTPException(status_code=404, detail=f"Source not found: {source_id}")

#         if not source.full_text or not source.full_text.strip():
#             raise HTTPException(
#                 status_code=400,
#                 detail="Source has no text content to generate a mind map from",
#             )

#         mind_map: Optional[Dict[str, Any]] = None

#         # ── Path 1: DOCX structured extraction (IR documents) ───────────────
#         file_path = source.asset.file_path if source.asset else None
#         if file_path and os.path.exists(file_path) and file_path.lower().endswith(('.docx', '.doc')):
#             try:
#                 logger.info(f"[MindMap] Trying DOCX structured extraction for {source_id}")
#                 mind_map = await asyncio.to_thread(_build_mindmap_from_docx, file_path)
#                 if mind_map:
#                     logger.success(f"[MindMap] DOCX extraction succeeded for {source_id}: {mind_map.get('label')}")
#                 else:
#                     logger.info(f"[MindMap] DOCX extraction returned None — falling back to LLM pipeline")
#             except Exception as e:
#                 logger.warning(f"[MindMap] DOCX extraction failed: {e} — falling back to LLM pipeline")
#                 mind_map = None

#         # ── Path 2: LLM pipeline fallback ───────────────────────────────────
#         if not mind_map:
#             orchestrator = get_orchestrator()
#             asyncio.create_task(_safe_produce(orchestrator, source_id))

#             logger.info(f"[MindMap] Running LLM pipeline for source_id={source_id}")
#             raw_result = await orchestrator.pipeline.generate_from_source_id(source_id)
#             logger.debug(f"[MindMap] LLM result type={type(raw_result)}, preview={str(raw_result)[:300]}")

#             if isinstance(raw_result, str):
#                 try:
#                     mind_map = _parse_llm_json(raw_result)
#                 except (json.JSONDecodeError, ValueError) as exc:
#                     logger.warning(f"[MindMap] JSON parse failed — using fallback. Error: {exc}")
#                     mind_map = _build_fallback_mind_map(raw_result, title=source.title or "Mind Map")
#             elif isinstance(raw_result, dict):
#                 mind_map = raw_result
#             else:
#                 raise ValueError(f"Mind map generation returned unsupported type: {type(raw_result)}")

#         # ── Validate ─────────────────────────────────────────────────────────
#         if not mind_map or not mind_map.get("label"):
#             mind_map = {
#                 "label": source.title or "Generated Mind Map",
#                 "children": [{"label": "Unable to generate structured mind map"}],
#             }

#         # Ensure clean JSON
#         mind_map = json.loads(json.dumps(mind_map))

#         logger.success(f"[MindMap] Completed for source_id={source_id}: {mind_map.get('label')}")
#         return MindMapResponse(mind_map=mind_map, source_id=source_id)

#     except HTTPException:
#         raise
#     except ValueError as e:
#         logger.error(f"Mind map validation error: {e}")
#         raise HTTPException(status_code=422, detail=str(e))
#     except Exception as e:
#         import traceback
#         logger.error(f"Mind map generation failed for source {source_id}: {type(e).__name__}: {e}")
#         logger.debug(f"Traceback:\n{traceback.format_exc()}")
#         raise HTTPException(status_code=500, detail=f"Mind map generation failed: {str(e)}")


# async def _safe_produce(orchestrator, source_id: str):
#     """Fire-and-forget Kafka job publish — errors are logged, never raised."""
#     try:
#         await orchestrator.produce_jobs([source_id])
#     except Exception as e:
#         logger.warning(f"Kafka produce skipped for {source_id}: {e}")


# async def start_kafka_consumer():
#     """Start the KafkaMindMapOrchestrator consumer as a background task."""
#     try:
#         orchestrator = get_orchestrator()
#         logger.info("Starting KafkaMindMapOrchestrator consumer...")
#         await orchestrator.start_consumer()
#     except Exception as e:
#         logger.warning(f"Kafka consumer could not start (Kafka may be unavailable): {e}")


# # ---------------------------------------------------------------------------
# # Images endpoint
# # ---------------------------------------------------------------------------

# class SourceImagesResponse(BaseModel):
#     images: List[str]  # base64-encoded PNG strings
#     source_id: str
#     count: int


# class NodeSummaryRequest(BaseModel):
#     node_name: str
#     root_subject: str


# class NodeSummaryResponse(BaseModel):
#     summary: str
#     node_name: str
#     root_subject: str


# # ---------------------------------------------------------------------------
# # Node Summary endpoint
# # ---------------------------------------------------------------------------

# @router.post("/sources/{source_id}/node-summary", response_model=NodeSummaryResponse)
# async def get_node_summary(source_id: str, request: NodeSummaryRequest):
#     """Generate a detailed summary for a specific mind map node using the source content."""
#     try:
#         logger.info(f"Node summary request: source_id={source_id!r}, node={request.node_name!r}")
#         source_id = _decode_source_id(source_id)

#         try:
#             source = await Source.get(source_id)
#         except Exception:
#             raise HTTPException(status_code=404, detail=f"Source not found: {source_id}")

#         if not source or not source.full_text or not source.full_text.strip():
#             raise HTTPException(status_code=400, detail="Source has no text content")

#         from open_notebook.ai.provision import provision_langchain_model
#         from langchain_core.messages import HumanMessage, SystemMessage

#         context_text = source.full_text[:12000]

#         system_prompt = (
#             "You are an expert analyst. Given source document content, provide a detailed, "
#             "well-structured summary about a specific topic as it relates to the main subject. "
#             "Be thorough, cite specific facts from the source, and organize your response clearly. "
#             "Do not add information not present in the source. "
#             "STRICT FORMATTING RULES: "
#             "- Do NOT use any markdown headers (no #, ##, ### etc). "
#             "- Do NOT use --- horizontal rules. "
#             "- Use **bold** only for section titles and key terms. "
#             "- Use plain numbered lists or bullet points for structure. "
#             "- Write in clean plain text paragraphs."
#         )

#         user_prompt = (
#             f"Discuss what these sources say about '{request.node_name}', "
#             f"in the larger context of '{request.root_subject}'.\n\n"
#             f"Source content:\n{context_text}"
#         )

#         messages = [
#             SystemMessage(content=system_prompt),
#             HumanMessage(content=user_prompt),
#         ]

#         llm = await provision_langchain_model(
#             user_prompt, None, "chat", max_tokens=2048, temperature=0.3
#         )

#         def _run_llm():
#             return llm.invoke(messages).content

#         summary = await asyncio.to_thread(_run_llm)

#         # Clean up output
#         summary = re.sub(r"<think>.*?</think>", "", summary, flags=re.DOTALL).strip()
#         summary = re.sub(r"^#{1,6}\s+(.+)$", r"**\1**", summary, flags=re.MULTILINE)
#         summary = re.sub(r"^-{3,}$", "", summary, flags=re.MULTILINE)
#         summary = re.sub(r"\n{3,}", "\n\n", summary).strip()

#         logger.info(f"Node summary generated for '{request.node_name}' in source {source_id}")
#         return NodeSummaryResponse(
#             summary=summary,
#             node_name=request.node_name,
#             root_subject=request.root_subject,
#         )

#     except HTTPException:
#         raise
#     except Exception as e:
#         logger.error(f"Node summary failed for source {source_id}: {e}")
#         raise HTTPException(status_code=500, detail=f"Node summary failed: {str(e)}")


# # ---------------------------------------------------------------------------
# # Source Summary endpoint
# # ---------------------------------------------------------------------------

# class SourceSummaryResponse(BaseModel):
#     summary: str
#     source_id: str


# @router.post("/sources/{source_id}/summary", response_model=SourceSummaryResponse)
# async def get_source_summary(source_id: str):
#     """Generate a full summary of the source document."""
#     try:
#         source_id = _decode_source_id(source_id)
#         logger.info(f"Source summary request: source_id={source_id!r}")

#         from open_notebook.graphs.summary import SummaryPipeline, SummaryTextProcessor, SummaryLLMService
#         from open_notebook.ai.provision import provision_langchain_model

#         llm = await provision_langchain_model("", None, "chat", max_tokens=4096, temperature=0.3)

#         pipeline = SummaryPipeline(
#             processor=SummaryTextProcessor(),
#             llm_service=SummaryLLMService(llm),
#         )

#         result = await pipeline.generate_from_source_id(source_id)

#         logger.info(f"Source summary generated for source {source_id}")
#         return SourceSummaryResponse(summary=result["summary"], source_id=source_id)

#     except HTTPException:
#         raise
#     except Exception as e:
#         logger.error(f"Source summary failed for source {source_id}: {e}")
#         raise HTTPException(status_code=500, detail=f"Source summary failed: {str(e)}")


# # ---------------------------------------------------------------------------
# # Image extraction helpers
# # ---------------------------------------------------------------------------

# def _extract_images_from_pdf(file_path: str) -> List[str]:
#     """Extract images from a PDF using embedded XObject images + page rendering."""
#     import io
#     import fitz  # PyMuPDF
#     from PIL import Image

#     images_b64: List[str] = []
#     seen_sizes: set = set()

#     doc = fitz.open(file_path)

#     for page_index in range(len(doc)):
#         page = doc[page_index]
#         embedded = page.get_images(full=True)
#         page_has_embedded = False

#         for img_info in embedded:
#             xref = img_info[0]
#             try:
#                 base_image = doc.extract_image(xref)
#                 img_bytes = base_image["image"]
#                 img = Image.open(io.BytesIO(img_bytes))
#                 if img.width < 80 or img.height < 80:
#                     continue
#                 buf = io.BytesIO()
#                 img.convert("RGB").save(buf, format="PNG")
#                 raw = buf.getvalue()
#                 key = (img.width, img.height, raw[:64])
#                 if key in seen_sizes:
#                     continue
#                 seen_sizes.add(key)
#                 images_b64.append(base64.b64encode(raw).decode("utf-8"))
#                 page_has_embedded = True
#             except Exception as e:
#                 logger.warning(f"Skipping embedded image xref={xref} on page {page_index}: {e}")

#         if not page_has_embedded:
#             try:
#                 mat = fitz.Matrix(150 / 72, 150 / 72)
#                 pix = page.get_pixmap(matrix=mat, alpha=False)
#                 img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

#                 arr = np.array(img)
#                 white_ratio = (arr > 240).all(axis=2).mean()
#                 if white_ratio > 0.92:
#                     continue

#                 buf = io.BytesIO()
#                 img.save(buf, format="PNG")
#                 raw = buf.getvalue()
#                 key = (img.width, img.height, raw[:64])
#                 if key not in seen_sizes:
#                     seen_sizes.add(key)
#                     images_b64.append(base64.b64encode(raw).decode("utf-8"))
#             except Exception as e:
#                 logger.warning(f"Page render failed for page {page_index}: {e}")

#     doc.close()
#     return images_b64


# def _extract_images_from_docx(file_path: str) -> List[str]:
#     """Extract embedded images from a .docx file."""
#     import io
#     from PIL import Image
#     from docx import Document

#     images_b64: List[str] = []
#     seen: set = set()

#     doc = Document(file_path)
#     for rel in doc.part.rels.values():
#         if "image" in rel.reltype:
#             try:
#                 img_bytes = rel.target_part.blob
#                 img = Image.open(io.BytesIO(img_bytes))
#                 if img.width < 80 or img.height < 80:
#                     continue
#                 buf = io.BytesIO()
#                 img.convert("RGB").save(buf, format="PNG")
#                 raw = buf.getvalue()
#                 key = raw[:128]
#                 if key in seen:
#                     continue
#                 seen.add(key)
#                 images_b64.append(base64.b64encode(raw).decode("utf-8"))
#             except Exception as e:
#                 logger.warning(f"Skipping docx image: {e}")

#     return images_b64


# @router.get("/sources/{source_id}/images", response_model=SourceImagesResponse)
# async def get_source_images(source_id: str):
#     """Extract images from a source's PDF/file and return them as base64 PNGs."""
#     try:
#         source_id = _decode_source_id(source_id)

#         try:
#             source = await Source.get(source_id)
#         except Exception:
#             raise HTTPException(status_code=404, detail=f"Source not found: {source_id}")

#         file_path = source.asset.file_path if source.asset else None
#         if not file_path or not os.path.exists(file_path):
#             return SourceImagesResponse(images=[], source_id=source_id, count=0)

#         ext = os.path.splitext(file_path)[1].lower()
#         images_b64: List[str] = []

#         if ext == ".pdf":
#             images_b64 = await asyncio.to_thread(_extract_images_from_pdf, file_path)
#         elif ext in (".docx", ".doc"):
#             images_b64 = await asyncio.to_thread(_extract_images_from_docx, file_path)
#         else:
#             try:
#                 import io
#                 from PIL import Image
#                 img = Image.open(file_path)
#                 buf = io.BytesIO()
#                 img.convert("RGB").save(buf, format="PNG")
#                 images_b64.append(base64.b64encode(buf.getvalue()).decode("utf-8"))
#             except Exception as e:
#                 logger.warning(f"Could not read file as image for {source_id}: {e}")

#         logger.info(f"Extracted {len(images_b64)} images from source {source_id}")
#         return SourceImagesResponse(images=images_b64, source_id=source_id, count=len(images_b64))

#     except HTTPException:
#         raise
#     except Exception as e:
#         logger.error(f"Image extraction failed for source {source_id}: {e}")
#         raise HTTPException(status_code=500, detail=f"Image extraction failed: {str(e)}")




import asyncio
import base64
import os
import json
import re
from typing import Any, Dict, List, Optional
from urllib.parse import unquote

import numpy as np
from fastapi import APIRouter, HTTPException
from loguru import logger
from pydantic import BaseModel

from open_notebook.domain.notebook import Source

router = APIRouter()

KAFKA_BOOTSTRAP_SERVERS = os.environ.get("KAFKA_BOOTSTRAP_SERVERS", "kafka:9093")


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _decode_source_id(source_id: str) -> str:
    return unquote(source_id)


def _parse_llm_json(raw: str) -> dict:
    logger.debug(f"[_parse_llm_json] Raw input (first 300 chars): {raw[:300]}")
    original_raw = raw
    cleaned = re.sub(r"<think>[\s\S]*?</think>", "", raw, flags=re.DOTALL).strip()
    cleaned = re.sub(r'^\s*(?:json\.loads\s*\(|json\s*\(|python\s*\()', '', cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r'\)\s*$', '', cleaned).strip()

    fence_match = re.search(r"```(?:json)?\s*([\s\S]*?)\s*```", cleaned, flags=re.DOTALL)
    if fence_match:
        fence_content = fence_match.group(1).strip()
        try:
            result = json.loads(fence_content)
            if isinstance(result, dict) and result.get("label"):
                print("++++++++++++++++++++++++++++++",result)
                return result
        except json.JSONDecodeError:
            pass

    candidates = []
    depth = 0
    start = -1
    in_string = False
    escape_next = False

    for i, char in enumerate(cleaned):
        if escape_next:
            escape_next = False
            continue
        if char == '\\':
            escape_next = True
            continue
        if char == '"' and depth >= 0:
            in_string = not in_string
            continue
        if in_string:
            continue
        if char == '{':
            if depth == 0:
                start = i
            depth += 1
        elif char == '}':
            depth -= 1
            if depth == 0 and start != -1:
                candidate = cleaned[start: i + 1]
                try:
                    parsed = json.loads(candidate)
                    if isinstance(parsed, dict) and parsed.get("label"):
                        candidates.append(parsed)
                except json.JSONDecodeError:
                    pass
                start = -1

    if candidates:
        best = sorted(candidates, key=lambda x: len(json.dumps(x)), reverse=True)[0]
        return best

    try:
        result = json.loads(cleaned)
        if isinstance(result, dict) and result.get("label"):
            return result
    except json.JSONDecodeError:
        pass

    logger.error(f"[_parse_llm_json] FAILED to parse. Raw (500 chars):\n{original_raw[:500]}\n...")
    raise ValueError(f"Could not extract valid JSON with 'label' key. Raw preview: {original_raw[:150]}...")


def _build_fallback_mind_map(text: str, title: str = "Mind Map") -> dict:
    sentences = [s.strip() for s in re.split(r'(?<=[.!?])\s+', text) if len(s.strip()) > 20]
    chunk_size = max(3, len(sentences) // 5)
    children = []
    for i in range(0, len(sentences), chunk_size):
        chunk = sentences[i:i + chunk_size]
        children.append({
            "label": f"Section {i // chunk_size + 1}",
            "children": [{"label": s[:120]} for s in chunk]
        })
    return {
        "label": title,
        "children": children if children else [{"label": "No structured content found"}]
    }


# ---------------------------------------------------------------------------
# FIXED: Mind Map builder constants & helpers
# ---------------------------------------------------------------------------

_EMPTY_VALS = {
    'nil', 'n/a', 'na', 'not applicable', 'not available',
    '-', '', 'none', 'n.a.', 'not known', 'unknown', '---',
    'not mentioned', 'not provided', 'not given',
}

_SKIP_PERSONAL_KEYS = {
    'previous involvements', 'fir no', 'police station', 'status of case',
    'status of accused', 'action taken', 'source country', 'route of smuggling',
    'carrier', 'recipient', 'repayment', 'visit to india', 'circumstances',
    'case registered', 'tattoo image', 'deformity image', 'interrogation report',
    'network details', 'details of hide outs', 'points for follow',
    'part', 'section', 'report', 'reference', 'sr no', 'serial', 'sl no',
}

_PERSONAL_PRIORITY = [
    'Name', 'Code Name', 'Alias', 'Date Of Birth', 'Age', 'Sex', 'Nationality',
    'Religion', 'Caste', 'Marital Status', 'Complexion', 'Height', 'Weight',
    'Build', 'Hair', 'Eyes', 'Descriptive Roll', 'Mark Of Identification',
    'Education', 'Occupation', 'Bad Habits', 'Present Address', 'Permanent Address',
    'Facebook Id', 'Email Id', 'Mobile No',
]

_CRIMINAL_KEYS = [
    'how he got involved', 'how she got involved', 'criminal career',
    'criminal history', 'modus operandi', 'expertise in criminal',
    'occupation before joining', 'bad habits', 'gang', 'weapon',
    'arms', 'explosive', 'firing', 'crime',
]

_LEGAL_KEYS = [
    'previous involvement', 'legal history', 'fir detail', 'case detail',
    'arrest detail', 'case registered', 'arrested by', 'police station',
    'under section', 'status of case', 'status of accused',
]

_MOVEMENT_KEYS = [
    'movement', 'hideout', 'hide out', 'places', 'location',
    'travel', 'stayed', 'residence', 'address',
]


def _is_empty_val(val: str) -> bool:
    return val.lower().strip() in _EMPTY_VALS


def _should_skip_personal_key(key: str) -> bool:
    kl = key.lower().strip()
    return any(s in kl for s in _SKIP_PERSONAL_KEYS)


def _is_real_name(text: str) -> bool:
    """Returns True if text looks like a real person name, not a sentence fragment."""
    if not text or len(text) < 3:
        return False
    if len(text) > 60:
        return False
    # Mid-sentence punctuation = narrative text
    if re.search(r'[.!?]\s+[A-Za-z]', text):
        return False
    # Year = date = narrative
    if re.search(r'\b\d{4}\b', text):
        return False
    # Common sentence words = narrative fragment
    sentence_indicators = [
        ' was ', ' were ', ' had ', ' has ', ' have ', ' got ', ' get ',
        ' the ', ' and ', ' but ', ' who ', ' which ', ' that ', ' with ',
        ' from ', ' near ', ' along ', ' after ', ' before ', ' during ',
        ' murder', ' arrest', ' case ', ' court', ' police', ' station',
        ' driving ', ' robbery', ' snatching', ' killed', ' shot',
        ' deceased', ' accused', ' victim',
    ]
    tl = text.lower()
    if any(si in tl for si in sentence_indicators):
        return False
    words = text.split()
    if not (1 <= len(words) <= 6):
        return False
    first_alpha = re.sub(r'[^a-zA-Z]', '', words[0])
    if not first_alpha or not first_alpha[0].isupper():
        return False
    return True


def _truncate_label(text: str, max_len: int = 300) -> str:
    """Truncate label cleanly at word boundary — generous limit to avoid cutting values."""
    text = re.sub(r'\s+', ' ', text).strip()
    if len(text) <= max_len:
        return text
    truncated = text[:max_len].rsplit(' ', 1)[0]
    return truncated.rstrip('.,;:') + '...'


def _leaf(text: str, max_len: int = 300) -> Dict:
    return {"label": _truncate_label(text, max_len)}


def _section(label: str, children: list) -> Dict:
    return {"label": label, "children": children}


def _fuzzy_match(sec_key: str, keywords: List[str]) -> bool:
    sk = sec_key.lower()
    return any(kw in sk for kw in keywords)


def _split_sentences(text: str, min_len: int = 60) -> List[str]:
    """Split text into complete sentences, filtering short/incomplete ones."""
    return [
        s.strip() for s in re.split(r'(?<=[.!?])\s+', text)
        if len(s.strip()) >= min_len
    ]


# ---------------------------------------------------------------------------
# FIXED: _build_mindmap_from_docx
# ---------------------------------------------------------------------------

def _build_mindmap_from_docx(file_path: str) -> Optional[Dict[str, Any]]:
    """
    Build a clean, structured mind map directly from DOCX tables and PART-IV narrative.
    Returns None if the file is not a valid IR docx.

    FIXED:
    - Personal: garbage/narrative values filtered
    - Family: only real names (no sentence fragments)
    - Labels: properly truncated, no mid-sentence cut-offs
    - Sections: fuzzy keyword matching
    - Sentences: min 60 chars = complete sentences only
    """
    from api.routers.sources import _extract_profile_from_docx, _extract_part_iv_from_docx

    profile = _extract_profile_from_docx(file_path)
    part4 = _extract_part_iv_from_docx(file_path)

    personal: dict = profile.get('personal', {})
    family: list = profile.get('family', [])
    associates: list = profile.get('associates', [])
    main_person: str = profile.get('main_person', '') or personal.get('Name', 'Unknown Subject')
    part4_sections: dict = part4.get('sections', {})

    if not personal and not part4_sections:
        return None

    children: list = []

    # Field buckets — separate criminal/legal/address from personal
    _CRIMINAL_FIELD_KEYS = (
        'modus operandi', 'occupation before joining crime', 'expertise in criminal',
        'main financer', 'details of code name', 'weapons details', 'weapon', 'arms',
        'telephone nos', 'phone  numbers', 'phone numbers', 'e mail', 'fb ids',
        'facebook or other', 'economic status',
    )
    _LEGAL_FIELD_KEYS = (
        'arrested by', 'place of arrest', 'fir no', 'case no', 'under section',
        'police station', 'status of case', 'previous involvements',
    )
    _ADDRESS_FIELD_KEYS = (
        'present address', 'permanent address', 'residential address during',
    )
    _SKIP_PERSONAL = (
        'details of close friends',
    )

    # ── 1. Personal Profile ───────────────────────────────────────────────
    personal_leaves = []
    criminal_from_fields = []
    legal_from_fields = []
    address_from_fields = []
    shown_keys = set()

    # Strict personal-only priority list
    _PERSONAL_PRIORITY_STRICT = [
        'Name', 'Code Name', 'Alias', 'Date Of Birth', 'Age', 'Sex', 'Nationality',
        'Religion', 'Caste/Tribe/Sect', 'Marital Status', 'Complexion', 'Height',
        'Weight', 'Descriptive Roll', 'Mark Of Identification', 'Place Of Birth',
        'Educational Qualification', 'Occupation', 'Bad Habits', 'Habits',
        'Social Status', 'Country Visited', 'Parentage',
        'Personal Documents I.E Passport Number / Details Of Any Photo I.D.',
    ]

    for field in _PERSONAL_PRIORITY_STRICT:
        for k, v in personal.items():
            k_clean = k.strip()
            if k_clean.lower() != field.lower().strip():
                continue
            if _is_empty_val(v) or _should_skip_personal_key(k_clean):
                break
            personal_leaves.append(_leaf(f"{k_clean}: {re.sub(r'  +', ' ', v).strip()}"))
            shown_keys.add(k_clean.lower())
            break

    # Remaining fields — bucket by type, skip noise
    _SKIP_PERSONAL_EXTRA = ('details of close friends',)
    for k, v in personal.items():
        k_clean = k.strip()
        kl = k_clean.lower()
        if kl in shown_keys or _is_empty_val(v) or _should_skip_personal_key(k_clean):
            continue
        if any(s in kl for s in _SKIP_PERSONAL_EXTRA):
            shown_keys.add(kl)
            continue
        v_clean = re.sub(r'  +', ' ', v).strip()
        leaf = _leaf(f"{k_clean}: {v_clean}")
        if any(ak in kl for ak in _ADDRESS_FIELD_KEYS):
            address_from_fields.append(leaf)
        elif any(lk in kl for lk in _LEGAL_FIELD_KEYS):
            legal_from_fields.append(leaf)
        elif any(ck in kl for ck in _CRIMINAL_FIELD_KEYS):
            criminal_from_fields.append(leaf)
        else:
            if len(v_clean) < 200 and not re.search(r'[.!?]\s+[A-Z]', v_clean):
                personal_leaves.append(leaf)
        shown_keys.add(kl)

    if personal_leaves:
        # ── Group personal fields into sub-categories ─────────────────────
        _IDENTITY_KEYS = {'name', 'code name', 'alias', 'date of birth', 'age',
                          'sex', 'nationality', 'religion', 'caste', 'marital status', 'parentage'}
        _PHYSICAL_KEYS = {'complexion', 'height', 'weight', 'build', 'hair', 'eyes',
                          'descriptive roll', 'mark of identification', 'place of birth'}
        _BACKGROUND_KEYS = {'educational qualification', 'occupation', 'bad habits',
                            'habits', 'social status', 'economic status', 'country visited',
                            'personal documents'}

        identity_leaves, physical_leaves, background_leaves, other_leaves = [], [], [], []
        for leaf in personal_leaves:
            label_lower = leaf['label'].split(':')[0].lower().strip()
            if any(k in label_lower for k in _IDENTITY_KEYS):
                identity_leaves.append(leaf)
            elif any(k in label_lower for k in _PHYSICAL_KEYS):
                physical_leaves.append(leaf)
            elif any(k in label_lower for k in _BACKGROUND_KEYS):
                background_leaves.append(leaf)
            else:
                other_leaves.append(leaf)

        personal_children = []
        if identity_leaves:
            personal_children.append({"label": "Identity", "children": identity_leaves})
        if physical_leaves:
            personal_children.append({"label": "Physical Description", "children": physical_leaves})
        if background_leaves:
            personal_children.append({"label": "Background", "children": background_leaves})
        if other_leaves:
            personal_children.append({"label": "Other Details", "children": other_leaves})

        children.append(_section("Personal Profile", personal_children))

    # ── 2. Family & Relationships ─────────────────────────────────────────
    family_leaves = []
    seen_family = set()

    for member in family:
        name = (member.get('name') or '').strip()
        relation = (member.get('relation') or '').title()
        details = (member.get('details') or '').strip()

        if not _is_real_name(name):
            continue
        if name.lower() in seen_family:
            continue
        seen_family.add(name.lower())

        label = f"{relation}: {name}" if relation else name
        if details:
            detail_short = details.split('|')[0].strip()
            if len(detail_short) < 80:
                label += f" ({detail_short})"

        family_leaves.append(_leaf(label))

    if family_leaves:
        children.append(_section("Family & Relationships", family_leaves))

    # ── 3. Criminal Career ────────────────────────────────────────────────
    criminal_leaves = list(criminal_from_fields)

    for sec_key, sec_val in part4_sections.items():
        if not _fuzzy_match(sec_key, _CRIMINAL_KEYS):
            continue
        sentences = _split_sentences(sec_val, min_len=60)
        if sentences:
            criminal_leaves.append(
                _section(sec_key, [_leaf(s, max_len=130) for s in sentences[:5]])
            )
        elif sec_val.strip() and len(sec_val.strip()) >= 30:
            criminal_leaves.append(_leaf(f"{sec_key}: {sec_val.strip()}", max_len=130))

    # Gang associates
    gang_leaves = []
    seen_gang = set()
    for assoc in associates:
        name = (assoc.get('name') or '').strip()
        rel = (assoc.get('relation') or '').lower()

        if not _is_real_name(name):
            continue
        if name.lower() in seen_gang:
            continue
        seen_gang.add(name.lower())

        if any(g in rel for g in ['gang', 'gangster', 'associate', 'group', 'friend']):
            details = (assoc.get('details') or '').strip()
            label = name
            if details:
                detail_short = details.split('|')[0].strip()
                if len(detail_short) < 60:
                    label += f" — {detail_short}"
            gang_leaves.append(_leaf(label, max_len=110))

    if gang_leaves:
        criminal_leaves.append(_section("Gang Associates", gang_leaves))

    if criminal_leaves:
        children.append(_section("Criminal Career", criminal_leaves))

    # ── 4. Legal History ──────────────────────────────────────────────────
    legal_leaves = list(legal_from_fields)

    for sec_key, sec_val in part4_sections.items():
        if not _fuzzy_match(sec_key, _LEGAL_KEYS):
            continue
        sentences = _split_sentences(sec_val, min_len=40)
        if sentences:
            legal_leaves.append(_section(sec_key, [_leaf(s) for s in sentences[:6]]))
        elif sec_val.strip():
            legal_leaves.append(_leaf(f"{sec_key}: {sec_val.strip()}"))

    if legal_leaves:
        children.append(_section("Legal History", legal_leaves))

    # ── 5. Movements & Hideouts ───────────────────────────────────────────
    movement_leaves = list(address_from_fields)

    for sec_key, sec_val in part4_sections.items():
        if not _fuzzy_match(sec_key, _MOVEMENT_KEYS):
            continue
        sentences = _split_sentences(sec_val, min_len=40)
        if sentences:
            movement_leaves.append(_section(sec_key, [_leaf(s) for s in sentences[:5]]))
        elif sec_val.strip():
            movement_leaves.append(_leaf(f"{sec_key}: {sec_val.strip()}"))

    if movement_leaves:
        children.append(_section("Movements & Hideouts", movement_leaves))

    # ── 6. Remaining PART-IV sections ────────────────────────────────────
    all_handled = _CRIMINAL_KEYS + _LEGAL_KEYS + _MOVEMENT_KEYS

    for sec_key, sec_val in part4_sections.items():
        if _fuzzy_match(sec_key, all_handled):
            continue
        if not sec_val.strip() or len(sec_val.strip()) < 30:
            continue
        sentences = _split_sentences(sec_val, min_len=60)
        if sentences:
            children.append(
                _section(sec_key, [_leaf(s, max_len=130) for s in sentences[:4]])
            )
        else:
            children.append(_leaf(f"{sec_key}: {sec_val.strip()}", max_len=130))

    if not children:
        return None

    return {
        "label": main_person,
        "children": children,
    }


# ---------------------------------------------------------------------------
# Orchestrator (module-level singleton)
# ---------------------------------------------------------------------------

_orchestrator: Optional[Any] = None


def _build_orchestrator():
    from langchain_ollama import ChatOllama
    from open_notebook.graphs.mind_map import (
        EasyOCRService,
        IntelligenceLLMService,
        KafkaMindMapOrchestrator,
        MindMapPipeline,
        TextProcessor,
    )

    ollama_url = os.environ.get("OLLAMA_BASE_URL", "http://host.docker.internal:11434")
    kafka_servers = os.environ.get("KAFKA_BOOTSTRAP_SERVERS", "kafka:9093")
    ollama_model = os.environ.get("OLLAMA_MODEL", os.environ.get("DEFAULT_CHAT_MODEL", "llama3:latest"))
    print(f"Using Ollama model: {ollama_model}")
    logger.info(f"Building MindMapPipeline — model: {ollama_model}, Ollama: {ollama_url}, Kafka: {kafka_servers}")

    llm = ChatOllama(
        model=ollama_model,
        temperature=0,
        base_url=ollama_url,
        num_ctx=8192,
        repeat_penalty=1.1,
    )
    ocr_service = EasyOCRService()
    text_processor = TextProcessor()
    llm_service = IntelligenceLLMService(llm)
    pipeline = MindMapPipeline(
        ocr_service=ocr_service,
        processor=text_processor,
        llm_service=llm_service,
    )
    orchestrator = KafkaMindMapOrchestrator(
        pipeline=pipeline,
        bootstrap_servers=kafka_servers,
    )
    logger.info("KafkaMindMapOrchestrator ready")
    return orchestrator


def get_orchestrator():
    global _orchestrator
    if _orchestrator is None:
        _orchestrator = _build_orchestrator()
    return _orchestrator


# ---------------------------------------------------------------------------
# Models
# ---------------------------------------------------------------------------

class MindMapRequest(BaseModel):
    model_name: str = "llama3:latest"
    temperature: float = 0.0


class MindMapResponse(BaseModel):
    mind_map: Dict[str, Any]
    source_id: str


# ---------------------------------------------------------------------------
# Mind Map endpoint
# ---------------------------------------------------------------------------

@router.post("/sources/{source_id}/mindmap", response_model=MindMapResponse)
async def generate_mind_map(source_id: str, request: MindMapRequest):
    """Generate a mind map from a source's content."""
    try:
        source_id = _decode_source_id(source_id)

        try:
            source = await Source.get(source_id)
        except Exception:
            raise HTTPException(status_code=404, detail=f"Source not found: {source_id}")

        if not source.full_text or not source.full_text.strip():
            raise HTTPException(
                status_code=400,
                detail="Source has no text content to generate a mind map from",
            )

        mind_map: Optional[Dict[str, Any]] = None

        # ── Path 1: DOCX structured extraction ──────────────────────────
        file_path = source.asset.file_path if source.asset else None
        if file_path and os.path.exists(file_path) and file_path.lower().endswith(('.docx', '.doc')):
            try:
                logger.info(f"[MindMap] Trying DOCX structured extraction for {source_id}")
                mind_map = await asyncio.to_thread(_build_mindmap_from_docx, file_path)
                if mind_map:
                    logger.success(f"[MindMap] DOCX extraction succeeded for {source_id}: {mind_map.get('label')}")
                else:
                    logger.info(f"[MindMap] DOCX extraction returned None — falling back to LLM pipeline")
            except Exception as e:
                logger.warning(f"[MindMap] DOCX extraction failed: {e} — falling back to LLM pipeline")
                mind_map = None

        # ── Path 2: LLM pipeline fallback ───────────────────────────────
        if not mind_map:
            orchestrator = get_orchestrator()
            asyncio.create_task(_safe_produce(orchestrator, source_id))

            logger.info(f"[MindMap] Running LLM pipeline for source_id={source_id}")
            raw_result = await orchestrator.pipeline.generate_from_source_id(source_id)
            print(f"Raw LLM result (first 300 chars): {raw_result}")
            logger.debug(f"[MindMap] LLM result type={type(raw_result)}, preview={str(raw_result)[:300]}")

            if isinstance(raw_result, str):
                try:
                    mind_map = _parse_llm_json(raw_result)
                except (json.JSONDecodeError, ValueError) as exc:
                    logger.warning(f"[MindMap] JSON parse failed — using fallback. Error: {exc}")
                    mind_map = _build_fallback_mind_map(raw_result, title=source.title or "Mind Map")
            elif isinstance(raw_result, dict):
                mind_map = raw_result
            else:
                raise ValueError(f"Mind map generation returned unsupported type: {type(raw_result)}")

        # ── Validate ─────────────────────────────────────────────────────
        if not mind_map or not mind_map.get("label"):
            mind_map = {
                "label": source.title or "Generated Mind Map",
                "children": [{"label": "Unable to generate structured mind map"}],
            }

        mind_map = json.loads(json.dumps(mind_map))

        logger.success(f"[MindMap] Completed for source_id={source_id}: {mind_map.get('label')}")
        return MindMapResponse(mind_map=mind_map, source_id=source_id)

    except HTTPException:
        raise
    except ValueError as e:
        logger.error(f"Mind map validation error: {e}")
        raise HTTPException(status_code=422, detail=str(e))
    except Exception as e:
        import traceback
        logger.error(f"Mind map generation failed for source {source_id}: {type(e).__name__}: {e}")
        logger.debug(f"Traceback:\n{traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"Mind map generation failed: {str(e)}")


async def _safe_produce(orchestrator, source_id: str):
    try:
        await orchestrator.produce_jobs([source_id])
    except Exception as e:
        logger.warning(f"Kafka produce skipped for {source_id}: {e}")


async def start_kafka_consumer():
    try:
        orchestrator = get_orchestrator()
        logger.info("Starting KafkaMindMapOrchestrator consumer...")
        await orchestrator.start_consumer()
    except Exception as e:
        logger.warning(f"Kafka consumer could not start (Kafka may be unavailable): {e}")


# ---------------------------------------------------------------------------
# Images endpoint
# ---------------------------------------------------------------------------

class SourceImagesResponse(BaseModel):
    images: List[str]
    source_id: str
    count: int


class NodeSummaryRequest(BaseModel):
    node_name: str
    root_subject: str


class NodeSummaryResponse(BaseModel):
    summary: str
    node_name: str
    root_subject: str


# ---------------------------------------------------------------------------
# Node Summary endpoint
# ---------------------------------------------------------------------------

@router.post("/sources/{source_id}/node-summary", response_model=NodeSummaryResponse)
async def get_node_summary(source_id: str, request: NodeSummaryRequest):
    try:
        logger.info(f"Node summary request: source_id={source_id!r}, node={request.node_name!r}")
        source_id = _decode_source_id(source_id)

        try:
            source = await Source.get(source_id)
        except Exception:
            raise HTTPException(status_code=404, detail=f"Source not found: {source_id}")

        if not source or not source.full_text or not source.full_text.strip():
            raise HTTPException(status_code=400, detail="Source has no text content")

        from open_notebook.ai.provision import provision_langchain_model
        from langchain_core.messages import HumanMessage, SystemMessage

        context_text = source.full_text[:12000]

        system_prompt = (
            "You are an expert analyst. Given source document content, provide a detailed, "
            "well-structured summary about a specific topic as it relates to the main subject. "
            "Be thorough, cite specific facts from the source, and organize your response clearly. "
            "Do not add information not present in the source. "
            "STRICT FORMATTING RULES: "
            "- Do NOT use any markdown headers (no #, ##, ### etc). "
            "- Do NOT use --- horizontal rules. "
            "- Use **bold** only for section titles and key terms. "
            "- Use plain numbered lists or bullet points for structure. "
            "- Write in clean plain text paragraphs."
        )

        user_prompt = (
            f"Discuss what these sources say about '{request.node_name}', "
            f"in the larger context of '{request.root_subject}'.\n\n"
            f"Source content:\n{context_text}"
        )

        messages = [
            SystemMessage(content=system_prompt),
            HumanMessage(content=user_prompt),
        ]

        llm = await provision_langchain_model(
            user_prompt, None, "chat", max_tokens=2048, temperature=0.3
        )

        def _run_llm():
            return llm.invoke(messages).content

        summary = await asyncio.to_thread(_run_llm)

        summary = re.sub(r"<think>.*?</think>", "", summary, flags=re.DOTALL).strip()
        summary = re.sub(r"^#{1,6}\s+(.+)$", r"**\1**", summary, flags=re.MULTILINE)
        summary = re.sub(r"^-{3,}$", "", summary, flags=re.MULTILINE)
        summary = re.sub(r"\n{3,}", "\n\n", summary).strip()

        logger.info(f"Node summary generated for '{request.node_name}' in source {source_id}")
        return NodeSummaryResponse(
            summary=summary,
            node_name=request.node_name,
            root_subject=request.root_subject,
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Node summary failed for source {source_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Node summary failed: {str(e)}")


# ---------------------------------------------------------------------------
# Source Summary endpoint
# ---------------------------------------------------------------------------

class SourceSummaryResponse(BaseModel):
    summary: str
    source_id: str


@router.post("/sources/{source_id}/summary", response_model=SourceSummaryResponse)
async def get_source_summary(source_id: str):
    try:
        source_id = _decode_source_id(source_id)
        logger.info(f"Source summary request: source_id={source_id!r}")

        from open_notebook.graphs.summary import SummaryPipeline, SummaryTextProcessor, SummaryLLMService
        from open_notebook.ai.provision import provision_langchain_model

        llm = await provision_langchain_model("", None, "chat", max_tokens=4096, temperature=0.3)

        pipeline = SummaryPipeline(
            processor=SummaryTextProcessor(),
            llm_service=SummaryLLMService(llm),
        )

        result = await pipeline.generate_from_source_id(source_id)

        logger.info(f"Source summary generated for source {source_id}")
        return SourceSummaryResponse(summary=result["summary"], source_id=source_id)

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Source summary failed for source {source_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Source summary failed: {str(e)}")


# ---------------------------------------------------------------------------
# Image extraction helpers
# ---------------------------------------------------------------------------

def _extract_images_from_pdf(file_path: str) -> List[str]:
    import io
    import fitz
    from PIL import Image

    images_b64: List[str] = []
    seen_sizes: set = set()

    doc = fitz.open(file_path)

    for page_index in range(len(doc)):
        page = doc[page_index]
        embedded = page.get_images(full=True)
        page_has_embedded = False

        for img_info in embedded:
            xref = img_info[0]
            try:
                base_image = doc.extract_image(xref)
                img_bytes = base_image["image"]
                img = Image.open(io.BytesIO(img_bytes))
                if img.width < 80 or img.height < 80:
                    continue
                buf = io.BytesIO()
                img.convert("RGB").save(buf, format="PNG")
                raw = buf.getvalue()
                key = (img.width, img.height, raw[:64])
                if key in seen_sizes:
                    continue
                seen_sizes.add(key)
                images_b64.append(base64.b64encode(raw).decode("utf-8"))
                page_has_embedded = True
            except Exception as e:
                logger.warning(f"Skipping embedded image xref={xref} on page {page_index}: {e}")

        if not page_has_embedded:
            try:
                mat = fitz.Matrix(150 / 72, 150 / 72)
                pix = page.get_pixmap(matrix=mat, alpha=False)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

                arr = np.array(img)
                white_ratio = (arr > 240).all(axis=2).mean()
                if white_ratio > 0.92:
                    continue

                buf = io.BytesIO()
                img.save(buf, format="PNG")
                raw = buf.getvalue()
                key = (img.width, img.height, raw[:64])
                if key not in seen_sizes:
                    seen_sizes.add(key)
                    images_b64.append(base64.b64encode(raw).decode("utf-8"))
            except Exception as e:
                logger.warning(f"Page render failed for page {page_index}: {e}")

    doc.close()
    return images_b64


def _extract_images_from_docx(file_path: str) -> List[str]:
    import io
    from PIL import Image
    from docx import Document

    images_b64: List[str] = []
    seen: set = set()

    doc = Document(file_path)
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                img_bytes = rel.target_part.blob
                img = Image.open(io.BytesIO(img_bytes))
                if img.width < 80 or img.height < 80:
                    continue
                buf = io.BytesIO()
                img.convert("RGB").save(buf, format="PNG")
                raw = buf.getvalue()
                key = raw[:128]
                if key in seen:
                    continue
                seen.add(key)
                images_b64.append(base64.b64encode(raw).decode("utf-8"))
            except Exception as e:
                logger.warning(f"Skipping docx image: {e}")

    return images_b64


@router.get("/sources/{source_id}/images", response_model=SourceImagesResponse)
async def get_source_images(source_id: str):
    try:
        source_id = _decode_source_id(source_id)

        try:
            source = await Source.get(source_id)
        except Exception:
            raise HTTPException(status_code=404, detail=f"Source not found: {source_id}")

        file_path = source.asset.file_path if source.asset else None
        if not file_path or not os.path.exists(file_path):
            return SourceImagesResponse(images=[], source_id=source_id, count=0)

        ext = os.path.splitext(file_path)[1].lower()
        images_b64: List[str] = []

        if ext == ".pdf":
            images_b64 = await asyncio.to_thread(_extract_images_from_pdf, file_path)
        elif ext in (".docx", ".doc"):
            images_b64 = await asyncio.to_thread(_extract_images_from_docx, file_path)
        else:
            try:
                import io
                from PIL import Image
                img = Image.open(file_path)
                buf = io.BytesIO()
                img.convert("RGB").save(buf, format="PNG")
                images_b64.append(base64.b64encode(buf.getvalue()).decode("utf-8"))
            except Exception as e:
                logger.warning(f"Could not read file as image for {source_id}: {e}")

        logger.info(f"Extracted {len(images_b64)} images from source {source_id}")
        return SourceImagesResponse(images=images_b64, source_id=source_id, count=len(images_b64))

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Image extraction failed for source {source_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Image extraction failed: {str(e)}")